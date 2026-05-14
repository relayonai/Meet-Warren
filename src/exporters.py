"""Multi-format exporters for generated newsletters and blog posts.

Public API:
- to_pdf(html_str) -> bytes
- to_docx(result_dict, kind) -> bytes              kind ∈ {'newsletter','blog'}
- to_markdown(result_dict, kind) -> str
- to_eml(html_str, text_str, subject) -> bytes

PDF goes via WeasyPrint (high-fidelity HTML→PDF). DOCX and Markdown are
built directly from the structured result dict so they get clean styling
rather than a noisy HTML-to-Word conversion. EML is a standard multipart
RFC 822 file you can double-click into Mail/Outlook.
"""
from __future__ import annotations

import io
import json
import logging
import os
from datetime import datetime, timezone
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.utils import formatdate, make_msgid
from typing import Optional

# WeasyPrint dynamically loads Pango/HarfBuzz/Cairo via dlopen. On macOS
# Homebrew installs these to /usr/local/lib (Intel) or /opt/homebrew/lib
# (Apple Silicon); the system loader does NOT search those paths by default,
# so we pre-seed DYLD_FALLBACK_LIBRARY_PATH before importing weasyprint.
for _libdir in ("/usr/local/lib", "/opt/homebrew/lib"):
    if os.path.isdir(_libdir):
        existing = os.environ.get("DYLD_FALLBACK_LIBRARY_PATH", "")
        if _libdir not in existing.split(":"):
            os.environ["DYLD_FALLBACK_LIBRARY_PATH"] = (
                _libdir if not existing else f"{_libdir}:{existing}"
            )

log = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def _is_md_table_para(text: str) -> bool:
    """True if a paragraph split from section content looks like a markdown table."""
    from .design_elements import is_markdown_table
    return is_markdown_table(text)


def _docx_add_md_table(doc, text: str):
    """Parse a markdown table chunk and add a python-docx Table to doc."""
    from docx.shared import RGBColor, Pt
    lines = [l.strip() for l in text.strip().split("\n") if l.strip()]
    headers = [c.strip() for c in lines[0].strip("|").split("|")]
    rows = [
        [c.strip() for c in l.strip("|").split("|")]
        for l in lines[2:]
        if l.strip()
    ]
    n_cols = len(headers)
    if not n_cols:
        return
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers[:n_cols]):
        hdr_cells[i].text = str(h)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    for ri, row in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, cell in enumerate(row[:n_cols]):
            row_cells[ci].text = str(cell)
    doc.add_paragraph()


def _docx_add_ve_table(doc, ve: dict):
    """Add a visual_elements table dict as a python-docx Table."""
    from docx.shared import RGBColor, Pt
    title   = ve.get("title", "")
    headers = ve.get("headers") or []
    rows    = ve.get("rows") or []
    if not headers and not rows:
        return
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x5A, 0x64, 0x78)
    n_cols = len(headers) or (len(rows[0]) if rows else 0)
    if not n_cols:
        return
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers[:n_cols]):
        hdr_cells[i].text = str(h)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    for ri, row in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, cell in enumerate(row[:n_cols]):
            row_cells[ci].text = str(cell)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def to_pdf(html_str: str, *, base_url: Optional[str] = None) -> bytes:
    """Render the HTML newsletter/blog to a PDF byte string via WeasyPrint."""
    try:
        from weasyprint import HTML
    except Exception as exc:
        raise RuntimeError(
            "WeasyPrint is not available. Install with `pip install weasyprint` and "
            "ensure Pango/Cairo are installed (`brew install pango` on macOS)."
        ) from exc
    return HTML(string=html_str, base_url=base_url or os.getcwd()).write_pdf()


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _docx_styles_init(doc):
    """Set document-wide font + paragraph defaults."""
    from docx.shared import Pt
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _docx_add_heading(doc, text: str, level: int = 1):
    from docx.shared import RGBColor, Pt
    h = doc.add_heading(text or "", level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)  # navy
        if level == 0:
            run.font.size = Pt(26)
        elif level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
    return h


def _docx_meta_line(doc, text: str):
    from docx.shared import RGBColor, Pt
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x5A, 0x64, 0x78)
    run.italic = True
    return p


def _docx_link(paragraph, url: str, text: str):
    """Insert a real hyperlink into a python-docx paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1A4F8B")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(color); rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    _ = RGBColor  # silence unused-import lint


def _docx_newsletter(result: dict) -> bytes:
    from docx import Document
    doc = Document()
    _docx_styles_init(doc)

    if result.get("edition_label"):
        _docx_meta_line(doc, result["edition_label"])
    _docx_add_heading(doc, result.get("subject_line", "UK Personal Finance Digest"), level=0)

    if result.get("intro"):
        doc.add_paragraph(result["intro"])

    pick = result.get("editor_pick") or {}
    if pick.get("title"):
        _docx_add_heading(doc, "★ Editor's Pick", level=2)
        p = doc.add_paragraph()
        _docx_link(p, pick.get("url", "#"), pick.get("title", ""))
        if pick.get("why"):
            doc.add_paragraph(pick["why"])

    for s in result.get("sections", []) or []:
        _docx_add_heading(doc, s.get("heading", ""), level=1)
        if s.get("summary"):
            _docx_meta_line(doc, s["summary"])
        for a in s.get("articles", []) or []:
            p = doc.add_paragraph(style="List Bullet")
            _docx_link(p, a.get("url", "#"), a.get("title", ""))
            if a.get("source"):
                p.add_run(f"  — {a['source']}").italic = True
            if a.get("blurb"):
                doc.add_paragraph(a["blurb"])
            if a.get("why_it_matters"):
                wm = doc.add_paragraph()
                wm.add_run("Why it matters: ").bold = True
                wm.add_run(a["why_it_matters"])
        if s.get("commentary"):
            cm = doc.add_paragraph()
            cm.add_run("Editor: ").bold = True
            cm.add_run(s["commentary"]).italic = True

    if result.get("closing"):
        doc.add_paragraph()
        doc.add_paragraph(result["closing"])
    if result.get("signature"):
        sig = doc.add_paragraph()
        sig.add_run(f"— {result['signature']}").italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_blog(result: dict) -> bytes:
    from docx import Document
    doc = Document()
    _docx_styles_init(doc)

    _docx_add_heading(doc, result.get("title", "Untitled"), level=0)
    if result.get("subtitle"):
        _docx_meta_line(doc, result["subtitle"])
    meta_bits = []
    if result.get("byline"):
        meta_bits.append(result["byline"])
    if result.get("reading_time_minutes"):
        meta_bits.append(f"{result['reading_time_minutes']} min read")
    if result.get("published_human"):
        meta_bits.append(result["published_human"])
    if meta_bits:
        _docx_meta_line(doc, "  ·  ".join(meta_bits))

    if result.get("key_takeaways"):
        _docx_add_heading(doc, "Key Takeaways", level=2)
        for t in result["key_takeaways"]:
            doc.add_paragraph(t, style="List Bullet")

    for para in (result.get("intro") or "").split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    # Build visual_elements lookup for section insertion
    _ve_by_section: dict = {}
    for ve in (result.get("visual_elements") or []):
        ai = ve.get("after_section")
        if ai is None:
            continue
        try:
            ai = int(ai)
        except (TypeError, ValueError):
            continue
        if ai >= 0:
            _ve_by_section.setdefault(ai, []).append(ve)

    for si, s in enumerate(result.get("sections", []) or []):
        _docx_add_heading(doc, s.get("heading", ""), level=1)
        for para in (s.get("content") or "").split("\n\n"):
            para = para.strip()
            if not para:
                continue
            if _is_md_table_para(para):
                _docx_add_md_table(doc, para)
            else:
                doc.add_paragraph(para)
        if s.get("pull_quote"):
            q = doc.add_paragraph()
            q.add_run(f'"{s["pull_quote"]}"').italic = True
        for ve in _ve_by_section.get(si, []):
            if ve.get("type") == "table":
                _docx_add_ve_table(doc, ve)
            elif (ve.get("type") or "").startswith("chart_"):
                p = doc.add_paragraph()
                p.add_run(f"[Chart: {ve.get('title', '')}]").italic = True

    _docx_add_heading(doc, "The Bottom Line", level=1)
    for para in (result.get("conclusion") or "").split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    if result.get("faqs"):
        _docx_add_heading(doc, "Frequently Asked Questions", level=1)
        for f in result["faqs"]:
            q = doc.add_paragraph()
            q.add_run(f"Q. {f.get('question','')}").bold = True
            doc.add_paragraph(f"A. {f.get('answer','')}")

    if result.get("sources_cited"):
        _docx_add_heading(doc, "Sources", level=2)
        for s in result["sources_cited"]:
            p = doc.add_paragraph(style="List Bullet")
            _docx_link(p, s.get("url", "#"), s.get("title", ""))
            if s.get("source"):
                p.add_run(f"  — {s['source']}").italic = True

    if result.get("seo_tags"):
        _docx_meta_line(doc, "Tags: " + ", ".join(f"#{t}" for t in result["seo_tags"]))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_docx(result: dict, *, kind: str) -> bytes:
    if kind == "newsletter":
        return _docx_newsletter(result)
    if kind == "blog":
        return _docx_blog(result)
    raise ValueError(f"Unknown kind: {kind}")


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------

def _md_newsletter(result: dict) -> str:
    out = []
    if result.get("edition_label"):
        out += [f"_{result['edition_label']}_", ""]
    out += [f"# {result.get('subject_line','UK Personal Finance Digest')}", ""]
    if result.get("preheader"):
        out += [f"> {result['preheader']}", ""]
    if result.get("intro"):
        out += [result["intro"], ""]

    pick = result.get("editor_pick") or {}
    if pick.get("title"):
        out += ["## ★ Editor's Pick",
                f"**[{pick.get('title','')}]({pick.get('url','#')})**", ""]
        if pick.get("why"):
            out += [pick["why"], ""]

    for s in result.get("sections", []) or []:
        out += [f"## {s.get('heading','')}", ""]
        if s.get("summary"):
            out += [f"_{s['summary']}_", ""]
        for a in s.get("articles", []) or []:
            line = f"- **[{a.get('title','')}]({a.get('url','#')})**"
            if a.get("source"):
                line += f" — _{a['source']}_"
            out.append(line)
            if a.get("blurb"):
                out.append(f"  {a['blurb']}")
            if a.get("why_it_matters"):
                out.append(f"  > **Why it matters:** {a['why_it_matters']}")
        if s.get("commentary"):
            out += ["", f"> _Editor: {s['commentary']}_", ""]
        out.append("")

    if result.get("closing"):
        out += ["---", "", result["closing"], ""]
    if result.get("signature"):
        out += [f"_— {result['signature']}_"]
    return "\n".join(out)


def _md_blog(result: dict) -> str:
    out: list[str] = []

    # YAML frontmatter — required for SEO/E-E-A-T tooling and recognised by
    # Hugo, Jekyll, Astro, MDX, and the AgriciDaniel/claude-blog analyser.
    def _yaml_str(value: str) -> str:
        # YAML-safe string: escape double quotes, then wrap in quotes.
        return '"' + (value or "").replace('"', '\\"') + '"'

    title       = result.get("title", "Untitled")
    description = result.get("meta_description") or result.get("subtitle") or ""
    author      = result.get("byline") or "The Warren Editorial Desk"
    published   = result.get("published_iso") or ""
    tags        = result.get("seo_tags") or []

    # slug — use the output basename if stamped in, otherwise slugify the title
    import re as _re
    slug = result.get("_output_basename") or ""
    if not slug:
        slug = _re.sub(r"[^a-z0-9]+", "-", title.lower()).strip("-")[:60]

    # primary keyword from first seo_tag
    keyword = tags[0] if tags else ""

    # og_image from env config
    og_image = ""
    try:
        from src.config import load_config as _load_cfg
        og_image = _load_cfg().warren_og_image
    except Exception:
        pass

    out_fm: list[str] = []
    out_fm.append(f"title: {_yaml_str(title)}")
    if description:
        out_fm.append(f"description: {_yaml_str(description)}")
    out_fm.append(f"author: {_yaml_str(author)}")
    if published:
        out_fm.append(f"date: {published}")
    if slug:
        out_fm.append(f"slug: {_yaml_str(slug)}")
    if keyword:
        out_fm.append(f"keyword: {_yaml_str(keyword)}")
    if og_image:
        out_fm.append(f"image: {_yaml_str(og_image)}")
        out_fm.append(f"og_image: {_yaml_str(og_image)}")
    if tags:
        out_fm.append(f"tags: [{', '.join(_yaml_str(t) for t in tags)}]")

    out.append("---")
    out += out_fm
    out.append("---")
    out.append("")

    out += [f"# {title}", ""]
    if result.get("subtitle"):
        out += [f"_{result['subtitle']}_", ""]
    meta_bits = []
    if result.get("byline"):                meta_bits.append(result["byline"])
    if result.get("reading_time_minutes"):  meta_bits.append(f"{result['reading_time_minutes']} min read")
    if result.get("published_human"):       meta_bits.append(result["published_human"])
    if meta_bits:
        out += ["_" + "  ·  ".join(meta_bits) + "_", ""]

    if result.get("key_takeaways"):
        out += ["## Key Takeaways"]
        out += [f"- {t}" for t in result["key_takeaways"]]
        out.append("")

    if result.get("intro"):
        out += [result["intro"], ""]

    _md_ve_by_section: dict = {}
    for ve in (result.get("visual_elements") or []):
        ai = ve.get("after_section")
        if ai is None:
            continue
        try:
            ai = int(ai)
        except (TypeError, ValueError):
            continue
        if ai >= 0:
            _md_ve_by_section.setdefault(ai, []).append(ve)

    for si, s in enumerate(result.get("sections", []) or []):
        out += [f"## {s.get('heading','')}", "", (s.get("content") or ""), ""]
        if s.get("pull_quote"):
            out += [f"> {s['pull_quote']}", ""]
        for ve in _md_ve_by_section.get(si, []):
            vtype = ve.get("type", "")
            title = ve.get("title", "")
            if vtype == "table":
                headers = ve.get("headers") or []
                rows = ve.get("rows") or []
                if headers:
                    out.append("| " + " | ".join(str(h) for h in headers) + " |")
                    out.append("| " + " | ".join("---" for _ in headers) + " |")
                    for row in rows:
                        out.append("| " + " | ".join(str(c) for c in row) + " |")
                    out.append("")
            elif vtype.startswith("chart_"):
                labels = ve.get("labels") or []
                values = ve.get("values") or []
                unit   = ve.get("unit", "")
                if title:
                    out.append(f"**{title}**")
                if labels and values:
                    out.append("| Label | Value |")
                    out.append("| --- | --- |")
                    for l, v in zip(labels, values):
                        out.append(f"| {l} | {v} {unit} |")
                    out.append("")

    _TRUST_FOOTER_MD = (
        "_The Warren Editorial Team produces independent UK personal finance analysis. "
        "[About Warren](https://meetwarren.co.uk/about) · "
        "[Contact](mailto:info@meetwarren.co.uk)_"
    )

    out += ["## The Bottom Line", "", result.get("conclusion", ""), "", _TRUST_FOOTER_MD, ""]

    if result.get("faqs"):
        out += ["## Frequently Asked Questions", ""]
        for f in result["faqs"]:
            out += [f"**Q. {f.get('question','')}**", "", f.get("answer", ""), ""]

    if result.get("sources_cited"):
        out += ["## Sources"]
        for s in result["sources_cited"]:
            out.append(f"- [{s.get('title','')}]({s.get('url','#')}) — _{s.get('source','')}_")
        out.append("")

    if result.get("seo_tags"):
        out.append("Tags: " + " ".join(f"#{t}" for t in result["seo_tags"]))

    # JSON-LD schema block at the bottom. Many CMS importers (Hugo, Astro, MDX)
    # accept raw HTML inline, and the quality analyser detects schema via either
    # BeautifulSoup OR a regex match on `"@type":` so the markdown gets credit
    # equal to the HTML output.
    try:
        from .blog_generator import build_jsonld
        schemas = build_jsonld(result)
        if schemas:
            out += ["", "<!-- JSON-LD schema below — recognised by Google + the blog quality analyser -->"]
            for d in schemas:
                out.append(
                    f'<script type="application/ld+json">'
                    f'{json.dumps(d, ensure_ascii=False)}'
                    f'</script>'
                )
    except Exception as e:
        log.warning("Could not embed JSON-LD in markdown: %s", e)

    return "\n".join(out)


def to_markdown(result: dict, *, kind: str) -> str:
    if kind == "newsletter":
        return _md_newsletter(result)
    if kind == "blog":
        return _md_blog(result)
    raise ValueError(f"Unknown kind: {kind}")


# ---------------------------------------------------------------------------
# EML (RFC 822 multipart, openable in Mail / Outlook / Thunderbird)
# ---------------------------------------------------------------------------

def to_eml(html_str: str, text_str: str, subject: str,
           *, sender: str = "Warren Editorial <noreply@meetwarren.co.uk>",
           recipient: str = "you@example.com") -> bytes:
    msg = MIMEMultipart("alternative")
    msg["Subject"] = subject
    msg["From"] = sender
    msg["To"] = recipient
    msg["Date"] = formatdate(localtime=True)
    msg["Message-ID"] = make_msgid(domain="meetwarren.co.uk")
    msg["X-Generated-At"] = datetime.now(timezone.utc).isoformat()
    msg.attach(MIMEText(text_str or "", "plain", "utf-8"))
    msg.attach(MIMEText(html_str or "", "html", "utf-8"))
    return msg.as_bytes()
