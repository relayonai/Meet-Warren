from __future__ import annotations

import json
import os
import subprocess
import sys
import tempfile
import threading
import time
import traceback
import uuid
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime, date

import dash
import dash_bootstrap_components as dbc
import pandas as pd
import plotly.express as px
from dash import Input, Output, State, ctx, dash_table, dcc, html, no_update

from src.answer_machine import (append_exemplar, delete_exemplar, draft_reply,
                                  load_knowledge_base)
from src.archive import (archive_stats, delete_entry as archive_delete_entry,
                          get_entry as archive_get_entry,
                          list_archive)
from src.blog_generator import generate_blog_post, blog_to_html, blog_to_text
from src.seo_agent import generate_seo_brief
from src.visual_extractor import extract_visuals
from src.blog_quality import quick_score as blog_quick_score
from src.compliance import ensure_compliant, load_rulebook
from src.compliance.advisor import advise_document, parse_elements as compliance_parse_elements
from src.compliance.pipeline import init_compliance_tables
from src.config import build_anthropic_client, load_config
from src.database import get_connection, init_db, query_articles
from src.exporters import to_pdf, to_docx, to_markdown, to_eml
from src.blog_quality_revision import revise_for_quality
from src.formatter import to_html, to_text
from src.generator import generate_newsletter
from src.internal_links import load_published_corpus
from src.brand_review import review_brand_voice
from src.source_verifier import summarise as verify_summarise, verify_urls

# ---------------------------------------------------------------------------
# Bootstrap
# ---------------------------------------------------------------------------

VENV_PYTHON = os.path.join(os.path.dirname(__file__), "venv", "bin", "python")
if not os.path.exists(VENV_PYTHON):
    VENV_PYTHON = sys.executable

cfg = load_config()

app = dash.Dash(
    __name__,
    external_stylesheets=[dbc.themes.FLATLY],
    title="Warren Workflow",
    suppress_callback_exceptions=True,
)

_scrape: dict = {"proc": None, "output_file": None}

FREQUENCIES = ["daily", "weekly", "monthly"]


# ---------------------------------------------------------------------------
# /downloads/<filename> — serves any file from the configured output_dir.
# Path traversal is blocked by abspath comparison.
# ---------------------------------------------------------------------------
from flask import abort, send_from_directory  # noqa: E402

@app.server.route("/downloads/<path:filename>")
def _serve_download(filename: str):
    safe_dir = os.path.abspath(cfg.output_dir)
    target   = os.path.abspath(os.path.join(safe_dir, filename))
    if not target.startswith(safe_dir + os.sep) or not os.path.isfile(target):
        abort(404)
    return send_from_directory(safe_dir, filename, as_attachment=True)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _get_df() -> pd.DataFrame:
    try:
        conn = get_connection(cfg.db_path)
        init_db(conn)
        rows = conn.execute(
            "SELECT id, title, url, source, published_at, created_at, "
            "category, relevance_score, scrape_frequency, summary FROM articles "
            "ORDER BY COALESCE(published_at, created_at) DESC"
        ).fetchall()
        conn.close()
        if not rows:
            return pd.DataFrame()
        records = []
        for r in rows:
            sd = {}
            if r["summary"]:
                try:
                    sd = json.loads(r["summary"])
                except Exception:
                    pass
            records.append({
                "id":               r["id"],
                "title":            r["title"],
                "url":              r["url"],
                "source":           r["source"],
                "published_at":     (r["published_at"] or r["created_at"] or "")[:10],
                "category":         r["category"] or "other",
                "relevance_score":  r["relevance_score"] or 0,
                "scrape_frequency": r["scrape_frequency"] or "daily",
                "summary_text":     sd.get("summary", ""),
                "key_points":       sd.get("key_points", []),
            })
        return pd.DataFrame(records)
    except Exception:
        return pd.DataFrame()


def _stat_card(title: str, value: str, color: str = "primary") -> html.Div:
    """Branded stat tile (replaces dbc.Card variant)."""
    return html.Div([
        html.Div(title, className="label"),
        html.Div(value, className="value"),
    ], className="warren-stat-card")


def _page_header(title: str, subtitle: str | None = None) -> html.Div:
    """Consistent page heading with the gold accent bar."""
    bits = [
        html.Div([
            html.Span(className="accent"),
            html.H3(title),
        ], className="page-title"),
    ]
    if subtitle:
        bits.append(html.P(subtitle, className="page-subtitle"))
    return html.Div(bits)


def _running_badge():
    return dbc.Badge("● Running", color="success", className="fs-6 p-2")

def _idle_badge():
    return dbc.Badge("○ Idle", color="secondary", className="fs-6 p-2")


# ---------------------------------------------------------------------------
# Sidebar + layout
# ---------------------------------------------------------------------------

SIDEBAR = html.Div([
    html.Div([
        html.Div([
            html.Span("W", className="warren-brand-logo"),
            html.Div([
                html.Div("Warren", className="warren-brand-text"),
                html.Div("Workflow", className="warren-brand-sub"),
            ]),
        ], className="warren-brand-mark"),
    ], className="warren-brand"),

    dbc.Nav([
        dbc.NavLink(
            [html.Span("🗄", className="me-2"), "Database"],
            href="/",
            active="exact",
            className="sidebar-link fw-semibold",
        ),
        dbc.NavLink(
            [html.Span("✍", className="me-2"), "Create"],
            href="/create",
            active="exact",
            className="sidebar-link fw-semibold",
        ),
        dbc.NavLink(
            [html.Span("🛡", className="me-2"), "Compliance"],
            href="/compliance",
            active="exact",
            className="sidebar-link fw-semibold",
        ),
        dbc.NavLink(
            [html.Span("💬", className="me-2"), "Answer Machine"],
            href="/answer-machine",
            active="exact",
            className="sidebar-link fw-semibold",
        ),
        dbc.NavLink(
            [html.Span("🗂", className="me-2"), "Archive"],
            href="/archive",
            active="exact",
            className="sidebar-link fw-semibold",
        ),
    ], vertical=True, pills=True, className="px-3 pt-3"),

    html.Div("UK Personal Finance · v1", className="warren-sidebar-footer"),
], className="warren-sidebar", style={
    "position": "fixed",
    "top": 0,
    "left": 0,
    "bottom": 0,
    "width": "230px",
    "zIndex": 100,
    "overflowY": "auto",
})

app.layout = html.Div([
    dcc.Location(id="url", refresh=False),
    SIDEBAR,
    html.Div(
        id="page-content",
        style={
            "marginLeft": "230px",
            "padding": "28px 36px 60px",
            "minHeight": "100vh",
        },
    ),
    # Shared stores / intervals
    dcc.Store(id="scrape-state", data={"running": False}),
    dcc.Interval(id="scrape-interval", interval=1000, disabled=True),
    dcc.Store(id="selected-article-ids", data=[]),
])

# ---------------------------------------------------------------------------
# Router
# ---------------------------------------------------------------------------

@app.callback(Output("page-content", "children"), Input("url", "pathname"))
def route(pathname):
    if pathname == "/create":
        return _create_page()
    if pathname == "/compliance":
        return _compliance_page()
    if pathname == "/answer-machine":
        return _answer_machine_page()
    if pathname == "/archive":
        return _archive_page()
    return _database_page()


# ===========================================================================
# DATABASE PAGE
# ===========================================================================

def _database_page():
    return html.Div([
        _page_header("Database",
                     "Articles fetched from RSS, GOV.UK and HTTP sources, deduped + scored."),
        dbc.Tabs([
            dbc.Tab(label="Overview",        tab_id="db-overview"),
            dbc.Tab(label="Article Browser", tab_id="db-browser"),
            dbc.Tab(label="Scrape",          tab_id="db-scrape"),
        ], id="db-tabs", active_tab="db-overview", className="mb-4"),
        html.Div(id="db-tab-content"),
    ])


@app.callback(Output("db-tab-content", "children"), Input("db-tabs", "active_tab"))
def render_db_tab(tab):
    if tab == "db-overview":
        return _overview_layout()
    if tab == "db-browser":
        return _browser_layout()
    if tab == "db-scrape":
        return _scrape_layout()
    return html.Div()


# --- Overview ---------------------------------------------------------------

def _overview_layout():
    return html.Div([
        dcc.Interval(id="overview-refresh", interval=30_000, n_intervals=0),
        html.Div(id="overview-content"),
    ])


@app.callback(Output("overview-content", "children"), Input("overview-refresh", "n_intervals"))
def refresh_overview(_):
    df = _get_df()
    if df.empty:
        return dbc.Alert("No articles yet — run a scrape first.", color="warning")

    total       = len(df)
    today_str   = date.today().isoformat()
    today_count = int((df["published_at"] == today_str).sum())
    sources     = df["source"].nunique()
    avg_score   = f"{df['relevance_score'].mean():.1f}"

    stat_row = dbc.Row([
        dbc.Col(_stat_card("Total Articles",  str(total),       "primary"), md=3),
        dbc.Col(_stat_card("Added Today",     str(today_count), "success"), md=3),
        dbc.Col(_stat_card("Active Sources",  str(sources),     "info"),    md=3),
        dbc.Col(_stat_card("Avg Relevance",   avg_score,        "warning"), md=3),
    ], className="mb-4 g-3")

    # By source
    src_df = df.groupby("source").size().reset_index(name="count").sort_values("count")
    fig_src = px.bar(src_df, x="count", y="source", orientation="h",
                     title="Articles by Source", color="count",
                     color_continuous_scale="Blues",
                     labels={"count": "Articles", "source": ""})
    fig_src.update_layout(coloraxis_showscale=False, plot_bgcolor="white",
                           margin=dict(l=10, r=10, t=40, b=10))

    # By category
    cat_df = df.groupby("category").size().reset_index(name="count")
    fig_cat = px.pie(cat_df, names="category", values="count",
                     title="Articles by Category", hole=0.4,
                     color_discrete_sequence=px.colors.qualitative.Set3)
    fig_cat.update_layout(margin=dict(l=10, r=10, t=40, b=10))

    # By frequency
    freq_df = df.groupby("scrape_frequency").size().reset_index(name="count")
    fig_freq = px.bar(freq_df, x="scrape_frequency", y="count",
                      title="Articles by Scrape Frequency",
                      color="scrape_frequency",
                      color_discrete_map={"daily": "#4CAF50", "weekly": "#2196F3", "monthly": "#FF9800"},
                      labels={"scrape_frequency": "Frequency", "count": "Articles"})
    fig_freq.update_layout(showlegend=False, plot_bgcolor="white",
                            margin=dict(l=10, r=10, t=40, b=10))

    # Relevance distribution
    fig_score = px.histogram(df, x="relevance_score", nbins=10,
                              title="Relevance Score Distribution",
                              color_discrete_sequence=["#2196F3"],
                              labels={"relevance_score": "Score"})
    fig_score.update_layout(plot_bgcolor="white", showlegend=False,
                             margin=dict(l=10, r=10, t=40, b=10))

    charts = dbc.Row([
        dbc.Col(dcc.Graph(figure=fig_src,   config={"displayModeBar": False}), md=6),
        dbc.Col(dcc.Graph(figure=fig_cat,   config={"displayModeBar": False}), md=6),
        dbc.Col(dcc.Graph(figure=fig_freq,  config={"displayModeBar": False}), md=6),
        dbc.Col(dcc.Graph(figure=fig_score, config={"displayModeBar": False}), md=6),
    ], className="g-3")

    return html.Div([stat_row, charts])


# --- Article Browser --------------------------------------------------------

def _browser_layout():
    df = _get_df()
    sources = sorted(df["source"].unique().tolist())    if not df.empty else []
    cats    = sorted(df["category"].unique().tolist())  if not df.empty else []

    return html.Div([
        dbc.Card(dbc.CardBody([
            dbc.Row([
                dbc.Col([
                    dbc.Label("Source"),
                    dcc.Dropdown(id="br-source", options=[{"label": s, "value": s} for s in sources],
                                 placeholder="All sources", multi=True, clearable=True),
                ], md=3),
                dbc.Col([
                    dbc.Label("Category"),
                    dcc.Dropdown(id="br-category", options=[{"label": c.title(), "value": c} for c in cats],
                                 placeholder="All categories", multi=True, clearable=True),
                ], md=3),
                dbc.Col([
                    dbc.Label("Frequency"),
                    dcc.Dropdown(id="br-frequency",
                                 options=[{"label": f.title(), "value": f} for f in FREQUENCIES],
                                 placeholder="All frequencies", multi=True, clearable=True),
                ], md=3),
                dbc.Col([
                    dbc.Label("Min Score"),
                    dcc.Slider(id="br-score", min=1, max=10, step=1, value=1,
                               marks={i: str(i) for i in range(1, 11)},
                               tooltip={"always_visible": False}),
                ], md=3),
            ], className="g-3"),
        ]), className="mb-3 shadow-sm"),

        dash_table.DataTable(
            id="br-table",
            columns=[
                {"name": "Title",     "id": "title",            "presentation": "markdown"},
                {"name": "Source",    "id": "source"},
                {"name": "Category",  "id": "category"},
                {"name": "Frequency", "id": "scrape_frequency"},
                {"name": "Score",     "id": "relevance_score"},
                {"name": "Date",      "id": "published_at"},
            ],
            data=[],
            row_selectable="single",
            selected_rows=[],
            page_size=15,
            sort_action="native",
            filter_action="native",
            style_table={"overflowX": "auto"},
            style_header={"backgroundColor": "#2196F3", "color": "white", "fontWeight": "bold"},
            style_data_conditional=[
                {"if": {"row_index": "odd"}, "backgroundColor": "#f9f9f9"},
                {"if": {"state": "selected"}, "backgroundColor": "#e3f2fd", "border": "1px solid #2196F3"},
                {"if": {"filter_query": '{scrape_frequency} = "daily"',   "column_id": "scrape_frequency"}, "color": "#2e7d32", "fontWeight": "600"},
                {"if": {"filter_query": '{scrape_frequency} = "weekly"',  "column_id": "scrape_frequency"}, "color": "#1565c0", "fontWeight": "600"},
                {"if": {"filter_query": '{scrape_frequency} = "monthly"', "column_id": "scrape_frequency"}, "color": "#e65100", "fontWeight": "600"},
            ],
            style_cell={"textAlign": "left", "padding": "10px", "fontSize": "13px"},
            style_cell_conditional=[
                {"if": {"column_id": "title"}, "maxWidth": "380px", "overflow": "hidden", "textOverflow": "ellipsis"},
                {"if": {"column_id": "relevance_score"}, "width": "60px", "textAlign": "center"},
            ],
        ),
        html.Div(id="br-detail", className="mt-3"),
    ])


@app.callback(
    Output("br-table", "data"),
    Input("br-source",    "value"),
    Input("br-category",  "value"),
    Input("br-frequency", "value"),
    Input("br-score",     "value"),
)
def update_browser(sources, cats, freqs, min_score):
    df = _get_df()
    if df.empty:
        return []
    if sources:
        df = df[df["source"].isin(sources)]
    if cats:
        df = df[df["category"].isin(cats)]
    if freqs:
        df = df[df["scrape_frequency"].isin(freqs)]
    if min_score:
        df = df[df["relevance_score"] >= min_score]
    display = df.copy()
    display["title"] = "[" + display["title"] + "](" + display["url"] + ")"
    return display[["title", "source", "category", "scrape_frequency", "relevance_score", "published_at"]].to_dict("records")


@app.callback(
    Output("br-detail", "children"),
    Input("br-table", "selected_rows"),
    State("br-table", "data"),
)
def show_br_detail(selected_rows, data):
    if not selected_rows or not data:
        return html.Div()
    row = data[selected_rows[0]]
    raw_title = row["title"].split("](")[0].lstrip("[")
    df = _get_df()
    match = df[df["title"].str.startswith(raw_title[:40])]
    if match.empty:
        return html.Div()
    art = match.iloc[0]
    kps = art.get("key_points") or []
    freq = art["scrape_frequency"]
    freq_color = {"daily": "success", "weekly": "primary", "monthly": "warning"}.get(freq, "secondary")
    return dbc.Card(dbc.CardBody([
        html.H5(raw_title, className="fw-bold"),
        dbc.Badge(art["category"].title(), color="primary",  className="me-2"),
        dbc.Badge(f"Score: {art['relevance_score']}", color="warning", text_color="dark", className="me-2"),
        dbc.Badge(freq.title(), color=freq_color, className="me-2"),
        html.Hr(),
        html.P(art["summary_text"] or "No summary available.", className="text-muted"),
        html.Ul([html.Li(kp) for kp in kps]) if kps else None,
    ]), className="shadow-sm", style={"borderLeft": "4px solid #2196F3"})


# --- Scrape -----------------------------------------------------------------

def _all_sources() -> list[tuple[str, str]]:
    """Return [(source_key, display_name)] for every configured source."""
    from src.scraper import RSS_SOURCE_OVERRIDES
    GOVUK_NAMES = {
        "office-for-national-statistics": "Office For National Statistics",
        "hm-revenue-customs":             "Hm Revenue Customs",
    }
    return (
        [(url, RSS_SOURCE_OVERRIDES.get(url, url)) for url in cfg.rss_feeds] +
        [(slug, GOVUK_NAMES.get(slug, slug)) for slug in cfg.govuk_orgs] +
        [(url, url) for url in cfg.http_sources]
    )


def _schedule_table() -> dbc.Table:
    """Build a 'last scraped' table from the DB. No schedule, manual runs only."""
    from src.database import get_source_log
    from src.scraper import RSS_SOURCE_OVERRIDES

    GOVUK_NAMES = {
        "office-for-national-statistics": "Office For National Statistics",
        "hm-revenue-customs":             "Hm Revenue Customs",
    }

    try:
        conn = get_connection(cfg.db_path)
        init_db(conn)
        log_rows = {r["source_key"]: r for r in get_source_log(conn)}
        conn.close()
    except Exception:
        log_rows = {}

    all_sources = (
        [(url, RSS_SOURCE_OVERRIDES.get(url, url)) for url in cfg.rss_feeds] +
        [(slug, GOVUK_NAMES.get(slug, slug)) for slug in cfg.govuk_orgs]
    )

    rows = []
    for key, name in all_sources:
        row    = log_rows.get(key)
        last   = row["last_scraped_at"][:19].replace("T", " ") if (row and row["last_scraped_at"]) else "Never"
        rows.append(html.Tr([
            html.Td(name,  style={"fontWeight": "500"}),
            html.Td(last,  className="text-muted", style={"fontSize": "13px"}),
        ]))

    return dbc.Table(
        [html.Thead(html.Tr([html.Th("Source"), html.Th("Last Scraped")])),
         html.Tbody(rows)],
        bordered=False, hover=True, size="sm", className="mb-0",
    )


def _scrape_layout():
    return html.Div([
        # Schedule status card
        dbc.Card(dbc.CardBody([
            dbc.Row([
                dbc.Col([
                    html.H6("Last Scraped", className="fw-bold text-primary mb-2"),
                    _schedule_table(),
                ], md=8),
                dbc.Col([
                    html.H6("Run Scrape", className="fw-bold text-primary mb-2"),
                    dbc.Button([
                        "▶  Scrape All Sources ",
                        html.Sup("ⓘ", style={"fontSize": "0.65em", "opacity": "0.7"}),
                    ], id="scrape-btn", color="primary", size="md", className="d-block w-100 mb-2"),
                    dbc.Tooltip("Manually scrapes every configured source right now. Scraping only runs when you trigger it.",
                                target="scrape-btn", placement="left"),

                    # Hidden legacy button — kept so existing callbacks remain valid.
                    dbc.Button("Force All", id="scrape-force-btn",
                               style={"display": "none"}, disabled=True),

                    dbc.Button([
                        "⏹  Stop ",
                        html.Sup("ⓘ", style={"fontSize": "0.65em", "opacity": "0.7"}),
                    ], id="scrape-stop-btn", color="danger", size="md", className="d-block w-100", disabled=True),
                    dbc.Tooltip("Terminates the running scrape process. Articles already summarised and stored will be kept.",
                                target="scrape-stop-btn", placement="left"),

                    html.Div(id="scrape-status-badge", className="mt-3 text-center"),
                ], md=4),
            ]),
        ]), className="mb-3 shadow-sm"),

        # On-demand source picker
        dbc.Card(dbc.CardBody([
            html.H6("On-Demand Scrape", className="fw-bold text-primary mb-2"),
            html.P("Pick any sources to scrape right now.",
                   className="text-muted small mb-3"),
            dbc.Row([
                dbc.Col([
                    dcc.Dropdown(
                        id="scrape-source-picker",
                        options=[{"label": name, "value": key} for key, name in _all_sources()],
                        multi=True,
                        placeholder="Select one or more sources…",
                    ),
                ], md=8),
                dbc.Col([
                    dbc.ButtonGroup([
                        dbc.Button("All", id="scrape-pick-all",  color="light", size="sm"),
                        dbc.Button("Clear", id="scrape-pick-clear", color="light", size="sm"),
                    ], className="me-2"),
                    dbc.Button([
                        "▶  Scrape Selected ",
                        html.Sup("ⓘ", style={"fontSize": "0.65em", "opacity": "0.7"}),
                    ], id="scrape-selected-btn", color="success", size="md", disabled=True),
                    dbc.Tooltip("Runs an immediate scrape against ONLY the sources you ticked above.",
                                target="scrape-selected-btn", placement="left"),
                ], md=4, className="text-end"),
            ]),
        ]), className="mb-3 shadow-sm"),

        # Live log
        dbc.Card(dbc.CardBody([
            html.H6("Live Output", className="text-muted mb-2"),
            html.Pre(id="scrape-log", children="No scrape run yet.", style={
                "backgroundColor": "#1e1e1e", "color": "#d4d4d4",
                "padding": "16px", "borderRadius": "6px",
                "height": "400px", "overflowY": "auto",
                "fontSize": "13px", "fontFamily": "monospace", "whiteSpace": "pre-wrap",
            }),
        ]), className="shadow-sm"),
    ])


@app.callback(
    Output("scrape-state",         "data"),
    Output("scrape-btn",           "disabled"),
    Output("scrape-force-btn",     "disabled"),
    Output("scrape-selected-btn",  "disabled", allow_duplicate=True),
    Output("scrape-stop-btn",      "disabled"),
    Output("scrape-interval",      "disabled"),
    Output("scrape-status-badge",  "children"),
    Input("scrape-btn",            "n_clicks"),
    Input("scrape-force-btn",      "n_clicks"),
    Input("scrape-selected-btn",   "n_clicks"),
    Input("scrape-stop-btn",       "n_clicks"),
    State("scrape-source-picker",  "value"),
    State("scrape-state",          "data"),
    prevent_initial_call=True,
)
def control_scrape(start_clicks, force_clicks, selected_clicks, stop_clicks,
                   selected_sources, state):
    trigger = ctx.triggered_id
    if trigger in ("scrape-btn", "scrape-force-btn", "scrape-selected-btn"):
        if _scrape.get("proc") and _scrape["proc"].poll() is None:
            return state, True, True, True, False, False, _running_badge()
        if trigger == "scrape-selected-btn" and not selected_sources:
            return state, False, False, True, True, True, _idle_badge()
        tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".log", delete=False)
        tmp.close()
        _scrape["output_file"] = tmp.name
        cmd = [VENV_PYTHON, "main.py", "scrape"]
        if trigger == "scrape-selected-btn":
            cmd += ["--sources", ",".join(selected_sources)]
        # scrape-btn / scrape-force-btn (legacy hidden) → scrape all sources, no flag
        _scrape["proc"] = subprocess.Popen(
            cmd,
            stdout=open(tmp.name, "w"),
            stderr=subprocess.STDOUT,
            cwd=os.path.dirname(os.path.abspath(__file__)),
        )
        return {"running": True}, True, True, True, False, False, _running_badge()
    if trigger == "scrape-stop-btn":
        if _scrape.get("proc") and _scrape["proc"].poll() is None:
            _scrape["proc"].terminate()
        return {"running": False}, False, False, False, True, True, _idle_badge()
    return state, False, False, False, True, True, _idle_badge()


@app.callback(
    Output("scrape-log",           "children"),
    Output("scrape-state",         "data",     allow_duplicate=True),
    Output("scrape-btn",           "disabled", allow_duplicate=True),
    Output("scrape-force-btn",     "disabled", allow_duplicate=True),
    Output("scrape-selected-btn",  "disabled", allow_duplicate=True),
    Output("scrape-stop-btn",      "disabled", allow_duplicate=True),
    Output("scrape-interval",      "disabled", allow_duplicate=True),
    Output("scrape-status-badge",  "children", allow_duplicate=True),
    Input("scrape-interval",       "n_intervals"),
    State("scrape-source-picker",  "value"),
    State("scrape-state",          "data"),
    prevent_initial_call=True,
)
def poll_scrape(_, picker_value, state):
    output_file = _scrape.get("output_file")
    proc        = _scrape.get("proc")
    sel_disabled_idle = not bool(picker_value)
    if not output_file or not proc:
        return no_update, state, False, False, sel_disabled_idle, True, True, _idle_badge()
    try:
        text = open(output_file).read()
    except Exception:
        text = ""
    if proc.poll() is None:
        return text or "Starting…", {"running": True}, True, True, True, False, False, _running_badge()
    return text or "(no output)", {"running": False}, False, False, sel_disabled_idle, True, True, _idle_badge()


@app.callback(
    Output("scrape-source-picker", "value"),
    Input("scrape-pick-all",   "n_clicks"),
    Input("scrape-pick-clear", "n_clicks"),
    prevent_initial_call=True,
)
def picker_quick_actions(_all, _clear):
    if ctx.triggered_id == "scrape-pick-all":
        return [k for k, _ in _all_sources()]
    return []


@app.callback(
    Output("scrape-selected-btn", "disabled"),
    Input("scrape-source-picker", "value"),
    State("scrape-state",         "data"),
)
def toggle_selected_btn(picker_value, state):
    if state and state.get("running"):
        return True
    return not bool(picker_value)


# ===========================================================================
# CREATE NEW CONTENT PAGE
# ===========================================================================

def _create_page():
    df = _get_df()
    sources = sorted(df["source"].unique().tolist())   if not df.empty else []
    cats    = sorted(df["category"].unique().tolist()) if not df.empty else []

    # ------------ LEFT PANE: filters + article picker ------------
    left_pane = html.Div([
        html.Div([
            html.Div("Step 1 · Filter & pick articles", className="section-eyebrow"),
            dbc.Row([
                dbc.Col([
                    dbc.Label("Source"),
                    dcc.Dropdown(id="cr-source",
                                 options=[{"label": s, "value": s} for s in sources],
                                 placeholder="All sources", multi=True, clearable=True),
                ], md=4),
                dbc.Col([
                    dbc.Label("Category"),
                    dcc.Dropdown(id="cr-category",
                                 options=[{"label": c.title(), "value": c} for c in cats],
                                 placeholder="All categories", multi=True, clearable=True),
                ], md=4),
                dbc.Col([
                    dbc.Label([
                        "Time frame ",
                        html.Sup("ⓘ", id="cr-timeframe-info",
                                 style={"fontSize": "0.7em", "opacity": "0.6", "cursor": "help"}),
                    ]),
                    dcc.Dropdown(id="cr-timeframe",
                                 options=[
                                     {"label": "Last 24h",  "value": "daily"},
                                     {"label": "Last 7d",   "value": "weekly"},
                                     {"label": "Last 30d",  "value": "monthly"},
                                     {"label": "All time",  "value": "all"},
                                 ],
                                 value="weekly", clearable=False),
                    dbc.Tooltip("Restricts the candidate pool to articles published within this window.",
                                target="cr-timeframe-info", placement="top"),
                ], md=4),
            ], className="g-3"),
            dbc.Row([
                dbc.Col([
                    dbc.Label("Min relevance score"),
                    dcc.Slider(id="cr-score", min=1, max=10, step=1, value=6,
                               marks={i: str(i) for i in range(1, 11)},
                               tooltip={"always_visible": False}),
                ], md=12),
            ], className="g-3 mt-1"),
            dbc.Row([
                dbc.Col([
                    dbc.Button("Select all (filtered)", id="cr-select-all",
                               color="outline-primary", size="sm", className="me-2"),
                    dbc.Button("Clear selection", id="cr-clear",
                               color="outline-secondary", size="sm"),
                ], md=12, className="mt-3"),
            ]),
        ], className="create-filter-strip"),

        html.Div(
            dash_table.DataTable(
                id="cr-table",
                columns=[
                    {"name": "",         "id": "select_hint",    "presentation": "markdown"},
                    {"name": "Title",    "id": "title"},
                    {"name": "Source",   "id": "source"},
                    {"name": "Category", "id": "category"},
                    {"name": "Score",    "id": "relevance_score"},
                    {"name": "Date",     "id": "published_at"},
                ],
                data=[],
                row_selectable="multi",
                selected_rows=[],
                page_size=14,
                sort_action="native",
                style_as_list_view=True,
                style_table={"overflowX": "auto"},
                style_data_conditional=[
                    {"if": {"row_index": "odd"},
                     "backgroundColor": "rgba(11,37,69,0.02)"},
                    {"if": {"state": "selected"},
                     "backgroundColor": "rgba(201,162,39,0.16)",
                     "border": "1px solid rgba(11,37,69,0.18)"},
                ],
                style_cell={"textAlign": "left", "padding": "10px 12px", "fontSize": "13px",
                            "fontFamily": "var(--font-body)", "color": "var(--warren-ink)"},
                style_header={"fontFamily": "var(--font-body)"},
                style_cell_conditional=[
                    {"if": {"column_id": "title"},  "maxWidth": "380px",
                     "overflow": "hidden", "textOverflow": "ellipsis", "fontWeight": "600"},
                    {"if": {"column_id": "select_hint"}, "width": "10px"},
                ],
            ),
            className="create-table-wrap",
        ),
    ], className="create-pane-left")

    # ------------ RIGHT PANE: type + sticky generate + output ------------
    right_pane = html.Div([
        # Selection summary chip
        html.Div(html.Div(id="cr-selected-badge"), className="mb-2"),

        # Step 2 — content type tiles
        dbc.Card(dbc.CardBody([
            html.Div("Step 2 · Content type", className="section-eyebrow"),
            html.Div([
                html.Div([
                    html.Div([
                        html.Div("✉", className="icon"),
                        html.Div("Newsletter", className="fw-bold mt-1"),
                        html.Div("Email digest with sections + editor pick.",
                                 className="text-muted small mt-1"),
                    ], style={"padding": "16px 14px", "textAlign": "center"}),
                ], id="cr-type-newsletter", n_clicks=0,
                   className="type-tile"),
                html.Div(style={"height": "10px"}),
                html.Div([
                    html.Div([
                        html.Div("📝", className="icon"),
                        html.Div("Blog Post", className="fw-bold mt-1"),
                        html.Div("Long-form analysis with TL;DR + FAQ.",
                                 className="text-muted small mt-1"),
                    ], style={"padding": "16px 14px", "textAlign": "center"}),
                ], id="cr-type-blog", n_clicks=0,
                   className="type-tile"),
            ]),
            dcc.Store(id="cr-content-type", data=None),
        ]), className="mb-3"),

        # Step 3 — optional editor's angle (priority framing)
        dbc.Card(dbc.CardBody([
            html.Div("Step 3 · Editor's angle (optional)", className="section-eyebrow"),
            dbc.Textarea(
                id="cr-angle",
                placeholder=("Optional: a 1–2 sentence brief telling the LLM how "
                             "to frame the piece.\n"
                             "e.g. 'Lead with the Bank of England rate cut and "
                             "what it means for first-time buyers.'"),
                style={"minHeight": "70px", "fontSize": "13px",
                       "fontFamily": "var(--font-body)"},
            ),
            html.Small(
                "Leave blank to let the LLM pick the angle from the article mix.",
                className="text-muted",
            ),
        ]), className="mb-3"),

        # Step 4 — sticky generate bar
        html.Div([
            html.Div("Step 4 · Generate", className="section-eyebrow",
                     style={"color": "rgba(255,255,255,0.7)"}),
            dbc.Button("⚡  Generate", id="cr-generate-btn", disabled=True),
            html.Div(id="cr-generate-hint", className="hint",
                     children="Pick at least one article and a content type."),
            dbc.Tooltip("Runs the LLM draft → compliance check → multi-format export. "
                        "Takes 30–90 seconds.",
                        target="cr-generate-btn", placement="top"),
        ], className="generate-bar"),

        # Job state + polling for progress
        dcc.Store(id="cr-job-id"),
        dcc.Interval(id="cr-job-poll", interval=600, disabled=True),
    ], className="create-pane-right")

    return html.Div([
        _page_header("Create",
                     "Filter articles, pick a format, ship a polished newsletter or blog."),
        html.Div([left_pane, right_pane], className="create-shell"),
        # Output sits below the two-pane shell so previews can use the full width
        html.Div(id="cr-output", className="mt-4"),
    ])


# --- Create page callbacks --------------------------------------------------

_TIMEFRAME_DAYS = {"daily": 1, "weekly": 7, "monthly": 30}


@app.callback(
    Output("cr-table", "data"),
    Input("cr-source",    "value"),
    Input("cr-category",  "value"),
    Input("cr-timeframe", "value"),
    Input("cr-score",     "value"),
)
def update_create_table(sources, cats, timeframe, min_score):
    df = _get_df()
    if df.empty:
        return []
    if sources:
        df = df[df["source"].isin(sources)]
    if cats:
        df = df[df["category"].isin(cats)]
    if timeframe and timeframe != "all":
        from datetime import datetime, timedelta
        cutoff = (datetime.utcnow() - timedelta(days=_TIMEFRAME_DAYS[timeframe])).strftime("%Y-%m-%d")
        df = df[df["published_at"] >= cutoff]
    if min_score:
        df = df[df["relevance_score"] >= min_score]
    display = df.copy()
    display["select_hint"] = ""
    return display[["select_hint", "title", "source", "category", "scrape_frequency",
                    "relevance_score", "published_at"]].to_dict("records")


@app.callback(
    Output("cr-table", "selected_rows", allow_duplicate=True),
    Input("cr-select-all", "n_clicks"),
    State("cr-table", "data"),
    prevent_initial_call=True,
)
def select_all(_, data):
    return list(range(len(data))) if data else []


@app.callback(
    Output("cr-table", "selected_rows", allow_duplicate=True),
    Input("cr-clear", "n_clicks"),
    prevent_initial_call=True,
)
def clear_selection(_):
    return []


@app.callback(
    Output("cr-selected-badge", "children"),
    Input("cr-table", "selected_rows"),
)
def update_selected_badge(selected):
    n = len(selected or [])
    if n == 0:
        return html.Span("No articles selected", className="text-muted small")
    return html.Span([
        html.Span("●", className="me-1"),
        f"{n} article{'s' if n != 1 else ''} selected",
    ], className="selection-chip")


@app.callback(
    Output("cr-type-newsletter", "className"),
    Output("cr-type-blog",       "className"),
    Output("cr-content-type",    "data"),
    Input("cr-type-newsletter",  "n_clicks"),
    Input("cr-type-blog",        "n_clicks"),
    State("cr-content-type",     "data"),
    prevent_initial_call=True,
)
def select_content_type(nl_clicks, blog_clicks, current):
    trigger = ctx.triggered_id
    if trigger == "cr-type-newsletter":
        return "type-tile selected", "type-tile", "newsletter"
    if trigger == "cr-type-blog":
        return "type-tile", "type-tile selected", "blog"
    return "type-tile", "type-tile", None


@app.callback(
    Output("cr-generate-btn",  "disabled", allow_duplicate=True),
    Output("cr-generate-hint", "children"),
    Input("cr-table",          "selected_rows"),
    Input("cr-content-type",   "data"),
    prevent_initial_call="initial_duplicate",
)
def update_generate_btn(selected, content_type):
    n = len(selected or [])
    if n == 0 and not content_type:
        return True, "Select at least one article and a content type first."
    if n == 0:
        return True, "Select at least one article."
    if not content_type:
        return True, "Choose a content type above."
    return False, f"Ready — {n} article{'s' if n != 1 else ''} selected as {content_type}."


# ---------------------------------------------------------------------------
# Generation jobs (background thread + polling) so the UI stays alive while
# Claude is drafting / compliance is running. Without this the whole callback
# blocks for 30–90s with no feedback.
# ---------------------------------------------------------------------------

_jobs: dict[str, dict] = {}        # job_id → {stage, label, started_at, error, result_div}
_jobs_lock = threading.Lock()

# Compliance-advisor jobs (separate dict, same pattern)
_cp_jobs: dict[str, dict] = {}
_cp_jobs_lock = threading.Lock()

_CP_STAGES = [
    ("parse",      "Parsing document elements"),
    ("hard_rules", "Checking banned phrases & terms"),
    ("llm",        "Running LLM compliance advisor"),
    ("done",       "Complete"),
]

# Stage labels shown in the progress card. Must match what the worker sets.
_STAGES = [
    ("collect",         "Loading article context"),
    ("seo_brief",       "SEO/AEO brief (Pass 1)"),
    ("draft",           "Drafting with Claude (Pass 2–3)"),
    ("visual_extract",  "Extracting visual elements (Pass 4)"),
    ("verify",          "Verifying source URLs"),
    ("quality_loop",    "Quality-revision loop (blog only)"),
    ("brand_review",    "Brand voice audit"),
    ("compliance",      "Compliance grading + revision"),
    ("export",          "Rendering PDF, DOCX, Markdown, EML"),
    ("quality",         "Final 100-pt rubric score"),
    ("done",            "Done"),
]


def _set_stage(job_id: str, stage_key: str, *, sub: str = "") -> None:
    with _jobs_lock:
        if job_id not in _jobs:
            return
        _jobs[job_id]["stage"] = stage_key
        _jobs[job_id]["sub"]   = sub


def _progress_card(job_id: str) -> html.Div:
    job = _jobs.get(job_id) or {}
    current_stage = job.get("stage", "collect")
    sub = job.get("sub", "")
    started = job.get("started_at") or time.time()
    elapsed = int(time.time() - started)

    # Stage list with check / spinner / dot icons.
    current_idx = next((i for i, (k, _) in enumerate(_STAGES) if k == current_stage), 0)
    rows = []
    for i, (key, label) in enumerate(_STAGES):
        if i < current_idx or key == "done" and current_stage == "done":
            icon = html.Span("✅", style={"width": "22px", "display": "inline-block"})
            cls  = "text-success fw-semibold"
        elif i == current_idx:
            icon = dbc.Spinner(size="sm", color="primary",
                               spinner_style={"width": "1rem", "height": "1rem",
                                              "marginRight": "4px"})
            cls  = "text-primary fw-bold"
        else:
            icon = html.Span("○", style={"width": "22px", "display": "inline-block",
                                          "color": "#c0c8d4"})
            cls  = "text-muted"
        sub_txt = (f"  —  {sub}" if (i == current_idx and sub) else "")
        rows.append(html.Div([icon, html.Span(label + sub_txt, className=cls)],
                              className="mb-2"))

    pct = int(((current_idx) / max(len(_STAGES) - 1, 1)) * 100)
    if current_stage == "done":
        pct = 100

    return html.Div([
        dbc.Card(dbc.CardBody([
            dbc.Row([
                dbc.Col([
                    html.H6("Generating your content", className="fw-bold mb-1"),
                    html.Small(f"Elapsed: {elapsed}s", className="text-muted"),
                ], md=8),
                dbc.Col([
                    dbc.Badge(f"{pct}%", color="primary", className="float-end fs-6"),
                ], md=4),
            ], className="mb-3"),
            dbc.Progress(value=pct, striped=True, animated=(current_stage != "done"),
                         className="mb-3", style={"height": "8px"}),
            html.Div(rows),
        ]), className="shadow-sm mb-3", style={"borderLeft": "4px solid #1a4f8b"}),
    ])


def _build_summaries(rows_df, raw_by_id) -> list[dict]:
    summaries = []
    for _, row in rows_df.iterrows():
        excerpt = (raw_by_id.get(row["id"], "") or "").strip()
        if len(excerpt) > 1200:
            excerpt = excerpt[:1200].rsplit(" ", 1)[0] + "…"
        summaries.append({
            "title":           row["title"],
            "url":             row["url"],
            "source":          row["source"],
            "published_at":    row["published_at"],
            "summary":         row["summary_text"],
            "key_points":      row["key_points"] if "key_points" in row else [],
            "category":        row["category"],
            "relevance_score": row["relevance_score"],
            "excerpt":         excerpt,
        })
    return summaries


def _versioned_basename(prefix: str) -> str:
    """e.g. 'newsletter-2026-04-25-v3' — picks next free version under output_dir."""
    today_str = datetime.utcnow().strftime("%Y-%m-%d")
    n = 1
    while True:
        base = f"{prefix}-{today_str}-v{n}"
        if not os.path.exists(os.path.join(cfg.output_dir, f"{base}.html")):
            return base
        n += 1


def _write_all_formats(base: str, *, kind: str, result: dict, html_str: str,
                       text_str: str, subject: str, summaries_in: list,
                       compliance_dict: dict) -> dict:
    """Write html/txt/json/md/pdf/docx/eml in parallel. Returns {ext: abs_path}."""
    paths = {
        "html": os.path.join(cfg.output_dir, f"{base}.html"),
        "txt":  os.path.join(cfg.output_dir, f"{base}.txt"),
        "md":   os.path.join(cfg.output_dir, f"{base}.md"),
        "pdf":  os.path.join(cfg.output_dir, f"{base}.pdf"),
        "docx": os.path.join(cfg.output_dir, f"{base}.docx"),
        "eml":  os.path.join(cfg.output_dir, f"{base}.eml"),
        "json": os.path.join(cfg.output_dir, f"{base}.json"),
    }

    def _write_html():  open(paths["html"], "w").write(html_str)
    def _write_txt():   open(paths["txt"],  "w").write(text_str)
    def _write_md():    open(paths["md"],   "w").write(to_markdown(result, kind=kind))
    def _write_pdf():   open(paths["pdf"], "wb").write(to_pdf(html_str))
    def _write_docx():  open(paths["docx"],"wb").write(to_docx(result, kind=kind))
    def _write_eml():   open(paths["eml"], "wb").write(to_eml(html_str, text_str, subject))
    def _write_json():
        with open(paths["json"], "w") as f:
            json.dump({
                "kind": kind,
                "generated_at_utc": datetime.utcnow().isoformat(),
                "subject_or_title": subject,
                "input_articles": summaries_in,
                "result": result,
                "compliance_summary": (compliance_dict or {})
                    .get("final_grade", {}).get("summary", {}),
            }, f, indent=2, ensure_ascii=False, default=str)

    jobs = {
        "html": _write_html, "txt": _write_txt,  "md":   _write_md,
        "pdf":  _write_pdf,  "docx": _write_docx, "eml":  _write_eml,
        "json": _write_json,
    }
    with ThreadPoolExecutor(max_workers=len(jobs)) as ex:
        futs = {ex.submit(fn): ext for ext, fn in jobs.items()}
        for fut, ext in futs.items():
            try:
                fut.result()
            except Exception as e:
                print(f"{ext.upper()} export failed: {e}")
                paths.pop(ext, None)
    return paths


def _run_generation_job(job_id: str, summaries: list, content_type: str,
                        *, editor_angle: str | None = None) -> None:
    """Heavy work: LLM draft → compliance → multi-format export. Runs in a thread."""
    try:
        client = build_anthropic_client(cfg)
        os.makedirs(cfg.output_dir, exist_ok=True)

        # --- Pass 1: SEO/AEO brief -------------------------------------------
        _set_stage(job_id, "seo_brief", sub="analysing keyword + AEO signals")
        seo_brief = None
        try:
            seo_brief = generate_seo_brief(
                summaries, client, cfg.anthropic_model,
                editor_angle=editor_angle,
            )
        except Exception as e:
            print(f"SEO brief failed (non-fatal): {e}")

        sub = f"{content_type} from {len(summaries)} article(s)"
        if editor_angle:
            sub += f" · angle: {editor_angle[:50]}…" if len(editor_angle) > 50 else f" · angle set"
        _set_stage(job_id, "draft", sub=sub)
        if content_type == "newsletter":
            result = generate_newsletter(summaries, client, cfg.anthropic_model,
                                          editor_angle=editor_angle,
                                          seo_brief=seo_brief)
            if not result:
                raise RuntimeError("Newsletter generation returned no content.")
            out_html, out_text = to_html(result), to_text(result)
            subject = result.get("subject_line", "UK Personal Finance Digest")
            base = _versioned_basename("newsletter")
            kind = "newsletter"
        elif content_type == "blog":
            # Pull prior Warren posts so the LLM can weave 3–5 internal links.
            # The post being generated doesn't exist on disk yet, so nothing
            # to exclude — but pass exclude_basename anyway for symmetry.
            existing_corpus = load_published_corpus(cfg.output_dir)
            result = generate_blog_post(
                summaries, client, cfg.anthropic_model,
                existing_posts=existing_corpus,
                editor_angle=editor_angle,
                seo_brief=seo_brief,
                progress_cb=lambda s: _set_stage(job_id, "draft", sub=s),
            )
            if not result:
                raise RuntimeError("Blog post generation returned no content.")
            out_html, out_text = blog_to_html(result), blog_to_text(result)
            subject = result.get("title", "Blog Post")
            base = _versioned_basename("blog")
            kind = "blog"
        else:
            raise RuntimeError(f"Unknown content type: {content_type}")

        # --- Pass 4: Visual extraction ---------------------------------------
        if result:
            _set_stage(job_id, "visual_extract",
                       sub=f"mining data from {kind} for visual elements")
            try:
                visual_elements = extract_visuals(
                    result, summaries, kind, client, cfg.anthropic_model,
                )
                if visual_elements:
                    result["visual_elements"] = visual_elements
            except Exception as e:
                print(f"Visual extraction failed (non-fatal): {e}")

        # --- Source verification (catches LLM-hallucinated URLs) ------------
        verification = None
        try:
            urls_to_check = []
            if kind == "blog":
                urls_to_check.extend(s.get("url", "")
                                     for s in (result.get("sources_cited") or []))
            if kind == "newsletter":
                ep = result.get("editor_pick") or {}
                if ep.get("url"):
                    urls_to_check.append(ep["url"])
                for s in result.get("sections", []) or []:
                    for a in s.get("articles", []) or []:
                        if a.get("url"):
                            urls_to_check.append(a["url"])
            urls_to_check = [u for u in urls_to_check if u]
            if urls_to_check:
                _set_stage(job_id, "verify",
                            sub=f"checking {len(urls_to_check)} URL(s)")
                records = verify_urls(urls_to_check, timeout=5)
                summary = verify_summarise(records)
                verification = {"records": records, "summary": summary,
                                 "checked": urls_to_check}
        except Exception as e:
            print(f"Source verification failed (non-fatal): {e}")
            verification = {"records": {}, "summary": {"total": 0, "ok": 0, "bad": 0,
                                                         "all_ok": True},
                             "checked": [], "error": str(e)}

        # --- Readability + flow pass (blog only) ----------------------------
        readability_result = None
        if kind == "blog":
            try:
                from src.readability_pass import run_readability_pass
                readability_result = run_readability_pass(
                    result, client=client, model=cfg.anthropic_model,
                    progress_cb=lambda s: _set_stage(job_id, "quality_loop", sub=s),
                )
                if readability_result.get("improved"):
                    result = readability_result["final_post"]
                    out_html = blog_to_html(result)
                    out_text = blog_to_text(result)
            except Exception as e:
                print(f"Readability pass failed (non-fatal): {e}")
                readability_result = {"improved": False, "error": str(e)}

        # --- Quality-revision loop (blog only) ------------------------------
        quality_revision = None
        if kind == "blog":
            _set_stage(job_id, "quality_loop", sub="scoring + revising weakest category")
            try:
                quality_revision = revise_for_quality(
                    result, client=client, model=cfg.anthropic_model, kind="blog",
                    progress_cb=lambda s: _set_stage(job_id, "quality_loop", sub=s),
                )
                if quality_revision.get("revised"):
                    # The loop produced a higher-scoring draft — adopt it.
                    result = quality_revision["final_post"]
                    out_html = blog_to_html(result)
                    out_text = blog_to_text(result)
            except Exception as e:
                print(f"Quality-revision loop failed (non-fatal): {e}")
                quality_revision = {"error": str(e), "revised": False, "audit": []}

            # Attach readability pass result to quality_revision for audit display
            if quality_revision and readability_result:
                quality_revision["readability_pass"] = {
                    "improved":           readability_result.get("improved", False),
                    "before_flesch":      readability_result.get("before_flesch"),
                    "after_flesch":       readability_result.get("after_flesch"),
                    "before_transitions": readability_result.get("before_transitions"),
                    "after_transitions":  readability_result.get("after_transitions"),
                    "changes":            readability_result.get("changes", []),
                }

        # --- Brand voice audit -----------------------------------------------
        brand_review = None
        try:
            _set_stage(job_id, "brand_review", sub="auditing voice + terminology")
            kb = load_knowledge_base()
            brand_review = review_brand_voice(
                out_html, kb=kb, client=client, model=cfg.anthropic_model,
                kind=kind,
            )
        except Exception as e:
            print(f"Brand review failed (non-fatal): {e}")
            brand_review = {"grade": "warn", "issues": [], "summary": f"Brand review failed: {e}",
                            "elapsed_seconds": 0.0, "error": str(e)}

        _set_stage(job_id, "compliance")
        conn_c = get_connection(cfg.db_path)
        compliance = ensure_compliant(
            out_html, kind=kind, client=client, model=cfg.anthropic_model,
            grader_model=cfg.compliance_model, conn=conn_c,
            content_ref=f"{base}.html",
            progress_cb=lambda s: _set_stage(job_id, "compliance", sub=s),
        )
        out_html = compliance["final_content"]
        # Re-render text from the (possibly revised) HTML? Source of truth is the
        # `result` dict, so plaintext stays in sync — only HTML is mutated.

        _set_stage(job_id, "export", sub="7 formats in parallel")
        result["_output_basename"] = base
        paths = _write_all_formats(base, kind=kind, result=result,
                                   html_str=out_html, text_str=out_text,
                                   subject=subject, summaries_in=summaries,
                                   compliance_dict=compliance)

        # --- Quality scoring (blog only — analyser is blog-tuned) -----------
        quality = None
        if kind == "blog":
            _set_stage(job_id, "quality", sub="100-pt rubric")
            try:
                # Score the markdown rendering (richer signal than HTML alone:
                # frontmatter + headings + links survive the markdown export).
                md_path = paths.get("md")
                if md_path and os.path.exists(md_path):
                    with open(md_path) as f:
                        md_text = f.read()
                else:
                    md_text = to_markdown(result, kind=kind)
                quality = blog_quick_score(md_text, suffix=".md")
            except Exception as e:
                print(f"Quality scoring failed: {e}")
                quality = {"error": str(e)}

        # Re-write the JSON sidecar with quality + revision + verification
        # blocks now that they're computed. The Archive page reads these.
        json_path = paths.get("json")
        if json_path and os.path.exists(json_path):
            try:
                with open(json_path) as f:
                    sidecar = json.load(f)
                if quality and "error" not in (quality or {}):
                    # Drop the heavy 'raw' analyser dump; the totals are enough.
                    q_slim = {k: v for k, v in quality.items() if k != "raw"}
                    sidecar["quality"] = q_slim
                if quality_revision is not None:
                    sidecar["quality_revision"] = {
                        "iterations":     quality_revision.get("iterations"),
                        "revised":        quality_revision.get("revised"),
                        "initial_score":  quality_revision.get("initial_score"),
                        "final_score":    (quality_revision.get("final_score") or {}).get("total"),
                        "audit":          quality_revision.get("audit"),
                    }
                if verification is not None:
                    sidecar["verification_summary"] = verification.get("summary")
                if brand_review is not None:
                    sidecar["brand_review"] = {
                        "grade":   brand_review.get("grade"),
                        "issues":  brand_review.get("issues", []),
                        "summary": brand_review.get("summary"),
                    }
                with open(json_path, "w") as f:
                    json.dump(sidecar, f, indent=2, ensure_ascii=False, default=str)
            except Exception as e:
                print(f"Could not update JSON sidecar with quality/revision data: {e}")

        # --- Build final preview component ----------------------------------
        if kind == "newsletter":
            sections_list = []
            if result.get("edition_label"):
                sections_list.append(dbc.ListGroupItem(f"🗓️ {result['edition_label']}"))
            pick = result.get("editor_pick") or {}
            if pick.get("title"):
                sections_list.append(dbc.ListGroupItem(f"★ Editor's Pick: {pick['title'][:60]}"))
            sections_list += [
                dbc.ListGroupItem(f"📌 {s.get('heading','')} — {len(s.get('articles',[]))} article(s)")
                for s in result.get("sections", [])
            ]
            preview = _content_preview("✉️ Newsletter Generated", subject,
                                        paths, out_html, sections_list, compliance,
                                        verification=verification,
                                        brand_review=brand_review)
        else:
            meta_list = []
            rt = result.get("reading_time_minutes")
            if rt:                          meta_list.append(dbc.ListGroupItem(f"⏱️ {rt} min read"))
            if result.get("byline"):        meta_list.append(dbc.ListGroupItem(f"✍️ {result['byline']}"))
            kt = result.get("key_takeaways") or []
            if kt:                          meta_list.append(dbc.ListGroupItem(f"💡 {len(kt)} key takeaways"))
            meta_list += [dbc.ListGroupItem(f"📌 {s.get('heading','')}")
                          for s in result.get("sections", [])]
            if result.get("faqs"):          meta_list.append(dbc.ListGroupItem(f"❓ {len(result['faqs'])} FAQs"))
            if result.get("sources_cited"): meta_list.append(dbc.ListGroupItem(f"🔗 {len(result['sources_cited'])} sources cited"))
            if result.get("seo_tags"):      meta_list.append(dbc.ListGroupItem("🏷️ " + " ".join(f"#{t}" for t in result["seo_tags"])))
            preview = _content_preview("📝 Blog Post Generated", subject,
                                        paths, out_html, meta_list, compliance,
                                        quality=quality, verification=verification,
                                        quality_revision=quality_revision,
                                        brand_review=brand_review)

        with _jobs_lock:
            _jobs[job_id]["stage"]      = "done"
            _jobs[job_id]["sub"]        = ""
            _jobs[job_id]["result_div"] = preview
    except Exception as e:
        traceback.print_exc()
        with _jobs_lock:
            _jobs[job_id]["stage"] = "done"
            _jobs[job_id]["error"] = f"{type(e).__name__}: {e}"


@app.callback(
    Output("cr-output",      "children", allow_duplicate=True),
    Output("cr-job-id",      "data"),
    Output("cr-job-poll",    "disabled"),
    Output("cr-generate-btn","disabled", allow_duplicate=True),
    Input("cr-generate-btn",  "n_clicks"),
    State("cr-table",         "selected_rows"),
    State("cr-table",         "data"),
    State("cr-content-type",  "data"),
    State("cr-angle",         "value"),
    prevent_initial_call=True,
)
def kickoff_generation(_, selected_rows, table_data, content_type, editor_angle):
    if not selected_rows or not table_data or not content_type:
        return (dbc.Alert("Nothing to generate — check your selections.",
                          color="warning"),
                no_update, True, False)

    selected_titles = [table_data[i]["title"] for i in selected_rows]
    df = _get_df()
    rows_df = df[df["title"].isin(selected_titles)]
    if rows_df.empty:
        return (dbc.Alert("Could not load article data.", color="danger"),
                no_update, True, False)

    # Pull raw_content excerpts so the editor synthesises real source material.
    try:
        conn_raw = get_connection(cfg.db_path)
        ids = [r["id"] for _, r in rows_df.iterrows()]
        placeholders = ",".join("?" * len(ids))
        raw_rows = conn_raw.execute(
            f"SELECT id, raw_content FROM articles WHERE id IN ({placeholders})", ids
        ).fetchall()
        conn_raw.close()
        raw_by_id = {r["id"]: (r["raw_content"] or "") for r in raw_rows}
    except Exception:
        raw_by_id = {}

    summaries = _build_summaries(rows_df, raw_by_id)
    job_id = uuid.uuid4().hex[:12]
    angle = (editor_angle or "").strip() or None
    with _jobs_lock:
        _jobs[job_id] = {
            "stage": "collect", "sub": "", "started_at": time.time(),
            "error": None, "result_div": None,
        }
    threading.Thread(target=_run_generation_job,
                     args=(job_id, summaries, content_type),
                     kwargs={"editor_angle": angle},
                     daemon=True).start()
    return _progress_card(job_id), job_id, False, True


@app.callback(
    Output("cr-output",       "children", allow_duplicate=True),
    Output("cr-job-poll",     "disabled", allow_duplicate=True),
    Output("cr-generate-btn", "disabled", allow_duplicate=True),
    Input("cr-job-poll",      "n_intervals"),
    State("cr-job-id",        "data"),
    prevent_initial_call=True,
)
def poll_generation(_n, job_id):
    if not job_id or job_id not in _jobs:
        return no_update, True, False
    job = _jobs[job_id]
    if job.get("error"):
        msg = dbc.Alert([html.Strong("Generation failed: "), job["error"]],
                        color="danger")
        # Free the slot
        _jobs.pop(job_id, None)
        return msg, True, False
    if job.get("stage") == "done" and job.get("result_div") is not None:
        result = job["result_div"]
        _jobs.pop(job_id, None)
        return result, True, False
    # Still running — refresh the progress card.
    return _progress_card(job_id), False, True



def _brand_review_card(brand_review: dict | None) -> html.Div:
    """Render a card summarising the brand voice audit result."""
    if not brand_review:
        return html.Div()

    grade   = brand_review.get("grade", "pass")
    issues  = brand_review.get("issues") or []
    summary = brand_review.get("summary", "")

    color = {"pass": "success", "warn": "warning", "fail": "danger"}.get(grade, "secondary")
    icon  = {"pass": "✅", "warn": "⚠️", "fail": "❌"}.get(grade, "•")

    criticals   = [i for i in issues if i.get("severity") == "critical"]
    warnings    = [i for i in issues if i.get("severity") == "warning"]
    suggestions = [i for i in issues if i.get("severity") == "suggestion"]

    _SEV_COLOR = {"critical": "danger", "warning": "warning", "suggestion": "info"}

    def _issue_item(issue: dict) -> html.Li:
        sev  = issue.get("severity", "suggestion")
        field = issue.get("field", "")
        finding = issue.get("finding", "")
        suggestion = issue.get("suggestion", "")
        return html.Li([
            dbc.Badge(sev.upper(), color=_SEV_COLOR.get(sev, "secondary"),
                      className="me-2", style={"fontSize": "0.7rem"}),
            dbc.Badge(field, color="light", text_color="dark",
                      className="me-2", style={"fontSize": "0.7rem"}),
            html.Strong(finding[:120]),
            html.Br(),
            html.Small(f"Fix: {suggestion}", className="text-muted"),
        ], style={"marginBottom": "10px"})

    issue_items = [_issue_item(i) for i in issues[:12]]

    stats_text = []
    if criticals:
        stats_text.append(f"{len(criticals)} critical")
    if warnings:
        stats_text.append(f"{len(warnings)} warning(s)")
    if suggestions:
        stats_text.append(f"{len(suggestions)} suggestion(s)")
    stats_str = ", ".join(stats_text) if stats_text else "No issues found"

    body_children = [
        dbc.Row([
            dbc.Col([
                html.Span(icon, style={"fontSize": "1.5rem"}),
                html.Span(" Brand Voice: ", className="ms-2"),
                dbc.Badge(grade.upper(), color=color, className="ms-1"),
                html.Span(f"  {stats_str}",
                          className="text-muted ms-2", style={"fontSize": "0.9rem"}),
            ]),
        ], className="mb-2"),
        html.P(summary, className="text-muted mb-0", style={"fontSize": "0.9rem"}),
    ]

    if issue_items:
        body_children.append(
            dbc.Accordion([
                dbc.AccordionItem(
                    html.Ul(issue_items, style={"paddingLeft": "1.2rem", "marginBottom": 0}),
                    title=f"View {len(issue_items)} brand voice issue(s)",
                ),
            ], start_collapsed=True, className="mt-3"),
        )

    return dbc.Card(
        dbc.CardBody(body_children),
        className="shadow-sm mb-3",
        style={"borderLeft": f"4px solid var(--bs-{color})"},
    )


def _compliance_card(compliance: dict) -> dbc.Card:
    """Render a card summarising the compliance result for the generated piece."""
    g       = (compliance or {}).get("final_grade", {}).get("summary", {})
    grade   = g.get("grade", "?")
    rate    = g.get("pass_rate", 0)
    failed  = g.get("failed", 0)
    total   = g.get("total", 0)
    iters   = compliance.get("iterations", 0)
    revised = compliance.get("revised", False)
    color   = {"pass": "success", "warn": "warning", "fail": "danger"}.get(grade, "secondary")
    icon    = {"pass": "✅", "warn": "⚠️", "fail": "❌"}.get(grade, "•")

    # Per-change rendering with §section pill extracted from the change string.
    import re as _re
    _SECTION_RE = _re.compile(r"§\s*([0-9]+(?:\.[0-9]+)*)")

    def _render_change(c: str) -> html.Span:
        m = _SECTION_RE.search(c or "")
        if m:
            sec = m.group(1)
            txt = _SECTION_RE.sub("", c).replace("()", "").strip(" .")
            return html.Span([
                dbc.Badge(f"§{sec}", color="warning", className="me-2",
                          style={"fontSize": "0.7rem"}),
                txt,
            ])
        return html.Span(c)

    grade_color = {"pass": "success", "warn": "warning", "fail": "danger"}
    audit_items = []
    for entry in compliance.get("audit", []) or []:
        i = entry.get("iteration", 0)
        g_i = entry.get("grade", "?")
        ch = entry.get("changes", []) or []
        if i == 0:
            audit_items.append(html.Li([
                html.Span("Initial grade: ", className="text-muted"),
                dbc.Badge(g_i.upper(), color=grade_color.get(g_i, "secondary"),
                          className="ms-1"),
            ], style={"marginBottom": "8px"}))
        else:
            audit_items.append(html.Li([
                html.Div([
                    html.Strong(f"Iteration {i} → "),
                    dbc.Badge(g_i.upper(), color=grade_color.get(g_i, "secondary"),
                              className="ms-1"),
                    html.Span(f"  ·  {len(ch)} change(s) applied",
                              className="text-muted ms-2", style={"fontSize": "0.85rem"}),
                ], className="mb-2"),
                html.Ul(
                    [html.Li(_render_change(c), style={"fontSize": "13px",
                                                       "marginBottom": "4px"})
                     for c in ch[:12]],
                    style={"marginTop": "4px", "marginBottom": "4px"},
                ),
            ], style={"marginBottom": "10px"}))

    failures = [e for e in compliance.get("final_grade", {}).get("expectations", [])
                if not e.get("passed")]
    fail_items = [
        html.Li([
            dbc.Badge(f"§{e.get('section','?')}", color="warning",
                      className="me-2", style={"fontSize": "0.7rem"}),
            html.Strong(e.get("text", "")),
            html.Br(),
            html.Small(e.get("evidence", "")[:160], className="text-muted"),
        ], style={"marginBottom": "10px"})
        for e in failures[:8]
    ]

    return dbc.Card(dbc.CardBody([
        dbc.Row([
            dbc.Col([
                html.Span(icon, style={"fontSize": "1.5rem"}),
                html.Span(" Compliance: ", className="ms-2"),
                dbc.Badge(grade.upper(), color=color, className="ms-1"),
                html.Span(f"  {int(rate*100)}% ({total - failed}/{total} checks passed)",
                          className="text-muted ms-2", style={"fontSize": "0.9rem"}),
            ], md=8),
            dbc.Col([
                dbc.Badge(f"Auto-revised x{iters}" if revised else "No revision needed",
                          color="info" if revised else "light", className="float-end"),
            ], md=4),
        ], className="mb-2"),
        dbc.Accordion([
            dbc.AccordionItem([
                html.H6("Audit trail", className="text-muted"),
                html.Ol(audit_items, style={"fontSize": "13px"}),
                html.H6("Outstanding issues", className="text-muted mt-3") if fail_items else None,
                html.Ul(fail_items) if fail_items else html.Em("None — all checks passed.",
                                                                className="text-success"),
            ], title="View compliance details"),
        ], start_collapsed=True, flush=True),
    ]), className="shadow-sm mb-3", style={"borderLeft": f"4px solid var(--bs-{color})"})


def _quality_revision_card(rev: dict | None) -> html.Div:
    """Render the audit trail of the quality-driven revision loop."""
    if not rev:
        return html.Div()
    if rev.get("error"):
        return dbc.Alert(f"Quality-revision loop errored: {rev['error']}",
                          color="warning", className="mb-3")
    audit = rev.get("audit", []) or []
    if not audit:
        return html.Div()

    initial = rev.get("initial_score") or 0
    final   = (rev.get("final_score") or {}).get("total", initial)
    iters   = rev.get("iterations", 0)
    revised = rev.get("revised", False)
    delta   = final - initial
    if delta > 0:
        delta_str = f"+{delta} pts"
        color = "success"
        icon = "📈"
    elif delta == 0:
        delta_str = "no change"
        color = "secondary"
        icon = "—"
    else:
        delta_str = f"{delta} pts"
        color = "warning"
        icon = "↘"

    audit_items = []
    for entry in audit:
        i = entry.get("iteration", 0)
        total = entry.get("total", 0)
        cat = entry.get("target_category")
        if i == 0:
            audit_items.append(html.Li([
                html.Strong("Initial draft "),
                dbc.Badge(f"{total}/100", color="secondary",
                          className="ms-1", style={"fontSize": "0.7rem"}),
            ], style={"marginBottom": "8px"}))
        else:
            improved = entry.get("improved", False)
            cat_label = {
                "content":   "Content Quality",
                "seo":       "SEO Optimization",
                "eeat":      "E-E-A-T Signals",
                "technical": "Technical Elements",
                "ai":        "AI Citation Readiness",
            }.get(cat, cat or "?")
            changes = entry.get("changes") or []
            audit_items.append(html.Li([
                html.Div([
                    html.Strong(f"Iteration {i} → "),
                    dbc.Badge(f"{total}/100",
                              color="success" if improved else "warning",
                              className="ms-1", style={"fontSize": "0.7rem"}),
                    html.Span(f"  ·  Target: {cat_label}",
                              className="text-muted ms-2",
                              style={"fontSize": "0.85rem"}),
                    html.Span(("  ·  ✓ adopted" if improved
                               else "  ·  ✗ no improvement, kept previous"),
                              className=("text-success" if improved else "text-warning"),
                              style={"fontSize": "0.8rem"}),
                ], className="mb-2"),
                html.Ul(
                    [html.Li(c, style={"fontSize": "13px",
                                        "marginBottom": "4px"})
                     for c in changes[:8]],
                    style={"marginTop": "4px"},
                ) if changes else html.Div(),
            ], style={"marginBottom": "10px"}))

    return dbc.Card(dbc.CardBody([
        dbc.Row([
            dbc.Col([
                html.Span(icon, style={"fontSize": "1.4rem"}),
                html.Span("  Quality revision: ", className="ms-2"),
                dbc.Badge(f"{initial} → {final}/100",
                          color=color, className="ms-1",
                          style={"fontSize": "0.85rem"}),
                html.Span(f"  ({delta_str})",
                          className="text-muted ms-1",
                          style={"fontSize": "0.85rem"}),
            ], md=8),
            dbc.Col([
                dbc.Badge(f"{iters} iteration{'s' if iters != 1 else ''}"
                          + ("" if revised else " · no rewrite"),
                          color="info" if revised else "light",
                          className="float-end"),
            ], md=4),
        ], className="mb-2"),
        dbc.Accordion([
            dbc.AccordionItem(
                html.Ol(audit_items, style={"fontSize": "13px"}),
                title="View revision audit trail",
            ),
        ], start_collapsed=True, flush=True),
    ]), className="shadow-sm mb-3",
        style={"borderLeft": f"4px solid var(--bs-{color})"})


def _verification_card(verification: dict | None) -> html.Div:
    """Render the source-URL verification result.

    Distinguishes three states per URL:
      ok      — 2xx/3xx response, verified live
      blocked — 401/403/429 from a paywalled/bot-protected site (likely valid)
      broken  — 404, 5xx, timeout, unreachable etc. (probably broken)

    Border colour reflects only `broken`. `blocked` is informational.
    """
    if not verification:
        return html.Div()
    summary = verification.get("summary", {}) or {}
    records = verification.get("records", {}) or {}
    if summary.get("total", 0) == 0:
        return html.Div()  # nothing was checked, no card

    total   = summary.get("total", 0)
    ok      = summary.get("ok", 0)
    blocked = summary.get("blocked", 0)
    broken  = summary.get("broken", 0)

    if broken == 0:
        color, icon = "success", "✅"
    elif broken / max(total, 1) <= 0.3:
        color, icon = "warning", "⚠️"
    else:
        color, icon = "danger", "❌"

    sev_color = {
        # Genuine failures
        "404": "danger", "4xx": "danger", "5xx": "danger",
        "timeout": "warning", "ssl": "warning", "unreachable": "warning",
        "invalid": "warning", "error": "warning",
        # Bot-blocked / paywalled — informational
        "blocked": "secondary",
    }

    def _row(u, r):
        status = r.get("status") or "?"
        is_blocked = status == "blocked"
        return html.Li([
            dbc.Badge(status.upper(),
                      color=sev_color.get(status, "secondary"),
                      className="me-2", style={"fontSize": "0.65rem"}),
            html.A(u, href=u, target="_blank",
                   style={"fontSize": "12px", "wordBreak": "break-all",
                          "color": "var(--warren-info)"}),
            html.Br(),
            html.Small(
                (r.get("note") if is_blocked
                 else f"HTTP {r.get('http_code')} · {r.get('error') or ''}"),
                className="text-muted",
            ),
        ], style={"marginBottom": "8px", "fontSize": "13px"})

    broken_records  = [(u, r) for u, r in records.items()
                        if r.get("status") in
                        {"404", "4xx", "5xx", "timeout", "ssl",
                         "unreachable", "invalid", "error"}]
    blocked_records = [(u, r) for u, r in records.items()
                        if r.get("status") == "blocked"]

    accordion_items = []
    if broken_records:
        accordion_items.append(dbc.AccordionItem([
            html.Ul([_row(u, r) for u, r in broken_records[:15]]),
            html.Small(
                ("Action: open these URLs and either fix the citation "
                 "(replace with a working source) or remove it before publishing."),
                className="text-muted"),
        ], title=f"⚠ {len(broken_records)} broken URL(s) — needs your attention"))
    if blocked_records:
        accordion_items.append(dbc.AccordionItem([
            html.Ul([_row(u, r) for u, r in blocked_records[:15]]),
            html.Small(
                ("These sites refused our automated check (paywall or bot "
                 "detection). The URL is almost certainly valid — open one "
                 "in your browser to spot-check if needed."),
                className="text-muted"),
        ], title=f"ℹ {len(blocked_records)} URL(s) blocked by paywall / bot detection"))

    if broken == 0 and blocked == 0:
        right_msg = html.Small("All cited URLs respond.",
                                className="text-success float-end")
    elif broken == 0:
        right_msg = html.Small(f"{blocked} blocked (paywall/bot detection)",
                                className="text-muted float-end")
    else:
        right_msg = html.Small(
            f"{broken} broken" + (f" · {blocked} blocked" if blocked else ""),
            className="text-danger float-end")

    return dbc.Card(dbc.CardBody([
        dbc.Row([
            dbc.Col([
                html.Span(icon, style={"fontSize": "1.4rem"}),
                html.Span("  Source verification: ", className="ms-2"),
                dbc.Badge(f"{ok} / {total} verified",
                          color=color, className="ms-1",
                          style={"fontSize": "0.85rem"}),
            ], md=8),
            dbc.Col([right_msg], md=4),
        ], className="mb-2"),

        (dbc.Accordion(accordion_items, start_collapsed=(broken == 0), flush=True)
         if accordion_items else html.Div()),
    ]), className="shadow-sm mb-3",
        style={"borderLeft": f"4px solid var(--bs-{color})"})


def _quality_card(quality: dict) -> dbc.Card:
    """Render the 100-pt blog quality breakdown."""
    if not quality:
        return html.Div()
    if quality.get("error"):
        return dbc.Alert(
            [html.Strong("Quality scoring failed: "), quality["error"]],
            color="warning", className="mb-3",
        )

    total = int(quality.get("total", 0))
    grade = quality.get("grade", "Unknown")
    cats  = quality.get("categories", {}) or {}
    maxs  = quality.get("max_per_category", {}) or {}
    issues = quality.get("top_issues", []) or []

    grade_colors = {
        "Exceptional":    "success",
        "Strong":         "success",
        "Acceptable":     "warning",
        "Below Standard": "warning",
        "Rewrite":        "danger",
    }
    color = grade_colors.get(grade, "secondary")
    accent_color = ("success" if total >= 80 else
                    "warning" if total >= 60 else "danger")

    cat_meta = [
        ("Content",      "content",   "Depth, readability, originality, structure, engagement, grammar"),
        ("SEO",          "seo",       "Title, headings, keywords, links, meta, URL"),
        ("E-E-A-T",      "eeat",      "Author, citations, trust, experience"),
        ("Technical",    "technical", "Schema, images, structured data, social"),
        ("AI Citation",  "ai",        "Citability, Q&A, entities, extraction"),
    ]
    cat_rows = []
    for label, key, blurb in cat_meta:
        v = int(cats.get(key, 0))
        m = int(maxs.get(key, 1)) or 1
        pct = int((v / m) * 100)
        bar_color = ("success" if pct >= 80 else
                     "warning" if pct >= 50 else "danger")
        cat_rows.append(dbc.Row([
            dbc.Col(html.Div([
                html.Strong(label, className="me-2"),
                html.Small(blurb, className="text-muted",
                           style={"fontSize": "11px"}),
            ]), md=6),
            dbc.Col(html.Div(f"{v} / {m}", className="text-end fw-bold"), md=2),
            dbc.Col(dbc.Progress(value=pct, color=bar_color,
                                 style={"height": "10px", "marginTop": "8px"}),
                    md=4),
        ], className="mb-2"))

    sev_color = {"high": "danger", "medium": "warning", "low": "secondary"}
    issue_items = [
        html.Li([
            dbc.Badge((iss.get("severity") or "low").upper(),
                      color=sev_color.get(iss.get("severity"), "secondary"),
                      className="me-2", style={"fontSize": "0.65rem"}),
            html.Span(f"[{iss.get('category', '')}] ", className="text-muted small"),
            iss.get("issue", ""),
        ], style={"marginBottom": "6px", "fontSize": "13px"})
        for iss in issues
    ]

    return dbc.Card(dbc.CardBody([
        dbc.Row([
            dbc.Col([
                html.Span("📊", style={"fontSize": "1.5rem"}),
                html.Span("  Blog Quality Score: ", className="ms-2"),
                dbc.Badge(f"{total}/100", color=accent_color,
                          className="ms-1", style={"fontSize": "0.95rem"}),
                html.Span(f"  · {grade}", className="text-muted ms-2",
                          style={"fontSize": "0.95rem"}),
                html.Div("AgriciDaniel/claude-blog · MIT-licensed analyzer",
                         className="text-muted",
                         style={"fontSize": "10px", "marginTop": "4px"}),
            ], md=8),
            dbc.Col([
                dbc.Badge(quality.get("content_type") or "post",
                          color="light", className="float-end"),
            ], md=4),
        ], className="mb-3"),

        html.Div(cat_rows, className="mb-3"),

        dbc.Accordion([
            dbc.AccordionItem([
                html.H6("Top quality fixes", className="text-muted"),
                html.Ul(issue_items) if issue_items
                  else html.Em("No issues flagged.", className="text-success"),
            ], title=f"View {len(issues)} prioritised fix(es)"),
        ], start_collapsed=True, flush=True),
    ]), className="shadow-sm mb-3",
        style={"borderLeft": f"4px solid var(--bs-{color})"})


_FORMAT_META = {
    "html": ("🌐 HTML",     "primary",  "Open / send as web page"),
    "pdf":  ("📄 PDF",      "danger",   "Print / share as PDF"),
    "docx": ("📝 Word",     "info",     "Edit in Microsoft Word"),
    "md":   ("⌨ Markdown", "dark",     "Paste into a CMS / Substack"),
    "eml":  ("✉ Email",    "warning",  "Open in Mail / Outlook to send"),
    "txt":  ("📃 Plain",    "secondary","Plain text fallback"),
    "json": ("{ } JSON",    "light",    "Raw structured data for replay"),
}
# Order shown in the UI
_FORMAT_ORDER = ["html", "pdf", "docx", "md", "eml", "txt", "json"]


def _download_bar(paths: dict) -> html.Div:
    buttons = []
    for ext in _FORMAT_ORDER:
        p = paths.get(ext)
        if not p:
            continue
        label, color, tip = _FORMAT_META[ext]
        fname = os.path.basename(p)
        buttons.append(dbc.Button(
            label, href=f"/downloads/{fname}", external_link=True,
            download=fname, color=color, size="sm",
            className="me-2 mb-2", title=tip,
        ))
    if not buttons:
        return html.Div()
    return html.Div([
        html.Div("Download as:", className="text-muted small mb-2"),
        html.Div(buttons),
    ], className="mb-3")


def _content_preview(badge_title: str, content_title: str,
                     paths: dict, preview_html: str, meta_items: list,
                     compliance: dict | None = None,
                     quality: dict | None = None,
                     verification: dict | None = None,
                     quality_revision: dict | None = None,
                     brand_review: dict | None = None) -> html.Div:
    primary = paths.get("html") or next(iter(paths.values()), "")
    return html.Div([
        dbc.Alert([
            html.Strong(f"✅ {badge_title}: {content_title}"),
            html.Br(),
            html.Small([
                "Saved to ",
                html.Code(os.path.dirname(primary), style={"fontSize": "0.78rem"}),
                f"  ({len(paths)} format(s))",
            ]),
        ], color="success", className="mb-3"),
        _download_bar(paths),
        _verification_card(verification) if verification else html.Div(),
        _quality_revision_card(quality_revision) if quality_revision else html.Div(),
        _brand_review_card(brand_review) if brand_review else html.Div(),
        _compliance_card(compliance) if compliance else html.Div(),
        _quality_card(quality) if quality else html.Div(),
        dbc.Row([
            dbc.Col([
                dbc.Card(dbc.CardBody([
                    html.H6("Contents", className="text-muted mb-2"),
                    dbc.ListGroup(meta_items, flush=True),
                ]), className="shadow-sm h-100"),
            ], md=3),
            dbc.Col([
                dbc.Card(dbc.CardBody([
                    dbc.Row([
                        dbc.Col(html.H6("Preview", className="text-muted mb-0 mt-1"), md=6),
                        dbc.Col([
                            dbc.ButtonGroup([
                                dbc.Button("📱 Mobile",  id="preview-mobile",  color="light",
                                           size="sm", outline=True, n_clicks=0),
                                dbc.Button("💻 Tablet",  id="preview-tablet",  color="light",
                                           size="sm", outline=True, n_clicks=0),
                                dbc.Button("🖥 Desktop", id="preview-desktop", color="primary",
                                           size="sm", n_clicks=0),
                            ], size="sm"),
                        ], md=6, className="text-end"),
                    ], className="mb-2"),
                    html.Div(
                        html.Iframe(
                            id="preview-frame",
                            srcDoc=preview_html,
                            style={"width": "100%", "height": "780px",
                                   "border": "1px solid #e6e9ef", "borderRadius": "6px",
                                   "background": "#ffffff"},
                        ),
                        id="preview-frame-wrap",
                        style={"width": "100%", "transition": "width 0.2s ease",
                               "margin": "0 auto"},
                    ),
                ]), className="shadow-sm"),
            ], md=9),
        ], className="g-3"),
    ])


@app.callback(
    Output("preview-frame-wrap", "style"),
    Output("preview-mobile",  "color"),
    Output("preview-mobile",  "outline"),
    Output("preview-tablet",  "color"),
    Output("preview-tablet",  "outline"),
    Output("preview-desktop", "color"),
    Output("preview-desktop", "outline"),
    Input("preview-mobile",  "n_clicks"),
    Input("preview-tablet",  "n_clicks"),
    Input("preview-desktop", "n_clicks"),
    prevent_initial_call=True,
)
def _toggle_preview_width(_m, _t, _d):
    trig = ctx.triggered_id
    base = {"transition": "width 0.2s ease", "margin": "0 auto"}
    if trig == "preview-mobile":
        return ({**base, "width": "390px"},
                "primary", False, "light", True, "light", True)
    if trig == "preview-tablet":
        return ({**base, "width": "768px"},
                "light", True, "primary", False, "light", True)
    return ({**base, "width": "100%"},
            "light", True, "light", True, "primary", False)


# ===========================================================================
# COMPLIANCE PAGE
# ===========================================================================

def _compliance_page():
    rb = load_rulebook()

    # ---- Recent compliance log -----------------------------------------------
    log_rows = []
    try:
        conn = get_connection(cfg.db_path)
        init_compliance_tables(conn)
        rows = conn.execute(
            "SELECT created_at, content_type, content_ref, grade, pass_rate, "
            "failed_count, iterations, revised "
            "FROM compliance_log ORDER BY id DESC LIMIT 50"
        ).fetchall()
        conn.close()
        log_rows = [dict(r) for r in rows]
    except Exception:
        pass

    # ---- Article compliance summary ------------------------------------------
    article_stats = {"pass": 0, "warn": 0, "fail": 0, "ungraded": 0}
    flagged_articles: list[dict] = []
    try:
        conn = get_connection(cfg.db_path)
        init_db(conn)
        rows = conn.execute(
            "SELECT title, source, compliance_grade, compliance_notes, published_at, created_at "
            "FROM articles ORDER BY COALESCE(published_at, created_at) DESC"
        ).fetchall()
        conn.close()
        for r in rows:
            g = r["compliance_grade"] or "ungraded"
            article_stats[g] = article_stats.get(g, 0) + 1
            if g in ("warn", "fail"):
                try:
                    notes = json.loads(r["compliance_notes"] or "[]")
                except Exception:
                    notes = []
                flagged_articles.append({
                    "title":  r["title"],
                    "source": r["source"],
                    "grade":  g,
                    "notes":  notes,
                    "date":   (r["published_at"] or r["created_at"] or "")[:10],
                })
    except Exception:
        pass

    # ---- UI ------------------------------------------------------------------
    summary_cards = dbc.Row([
        dbc.Col(_stat_card("Hard Rules",     str(len(rb.hard_rules)),         "primary"),  md=3),
        dbc.Col(_stat_card("Principles",     str(len(rb.principles)),          "info"),     md=3),
        dbc.Col(_stat_card("Articles Pass",  str(article_stats.get("pass", 0)),"success"), md=3),
        dbc.Col(_stat_card("Articles Flagged",
                           str(article_stats.get("warn", 0) + article_stats.get("fail", 0)),
                           "warning"), md=3),
    ], className="g-3 mb-4")

    log_table = dbc.Table([
        html.Thead(html.Tr([html.Th(h) for h in
            ["When", "Type", "Reference", "Grade", "Pass Rate", "Failed", "Iters", "Revised"]])),
        html.Tbody([
            html.Tr([
                html.Td(r["created_at"][:19].replace("T", " ")),
                html.Td(dbc.Badge(r["content_type"], color="secondary")),
                html.Td(html.Code(r["content_ref"] or "—", style={"fontSize": "0.78rem"})),
                html.Td(dbc.Badge(r["grade"].upper(),
                                  color={"pass": "success", "warn": "warning",
                                         "fail": "danger"}.get(r["grade"], "secondary"))),
                html.Td(f"{int((r['pass_rate'] or 0)*100)}%"),
                html.Td(str(r["failed_count"] or 0)),
                html.Td(str(r["iterations"] or 0)),
                html.Td("Yes" if r["revised"] else "No"),
            ]) for r in log_rows
        ]) if log_rows else html.Tbody(html.Tr(html.Td("No compliance runs yet.",
                                       colSpan=8, className="text-muted text-center p-3")))
    ], hover=True, size="sm", className="mb-3")

    flagged_table = dbc.Table([
        html.Thead(html.Tr([html.Th(h) for h in ["Article", "Source", "Date", "Grade", "Issues"]])),
        html.Tbody([
            html.Tr([
                html.Td(a["title"][:80]),
                html.Td(a["source"], className="text-muted small"),
                html.Td(a["date"], className="text-muted small"),
                html.Td(dbc.Badge(a["grade"].upper(),
                                  color={"warn": "warning", "fail": "danger"}.get(a["grade"], "secondary"))),
                html.Td(html.Ul([html.Li(n, style={"fontSize": "12px"}) for n in (a["notes"] or [])[:3]],
                                style={"marginBottom": 0, "paddingLeft": "16px"})),
            ]) for a in flagged_articles[:30]
        ]) if flagged_articles else html.Tbody(html.Tr(html.Td("No flagged articles.",
                                       colSpan=5, className="text-muted text-center p-3")))
    ], hover=True, size="sm")

    rules_html = dbc.Accordion([
        dbc.AccordionItem([
            html.H6("Banned phrases", className="text-muted mt-2"),
            html.Ul([html.Li([html.Code(r.pattern), f"  — §{r.section}: {r.rationale}"])
                     for r in rb.hard_rules if r.kind == "banned_phrase"]),
            html.H6("Banned terms", className="text-muted mt-3"),
            html.Ul([html.Li([html.Code(r.pattern), f"  — §{r.section}: {r.rationale}"])
                     for r in rb.hard_rules if r.kind == "banned_term"]),
            html.H6("Banned topics", className="text-muted mt-3"),
            html.Ul([html.Li([html.Code(r.pattern), f"  — §{r.section}: {r.rationale}"])
                     for r in rb.hard_rules if r.kind == "banned_topic"]),
        ], title=f"Hard Rules ({len(rb.hard_rules)})"),
        dbc.AccordionItem([
            html.Ul([html.Li([html.Strong(p.title), f"  (§{p.section})", html.Br(),
                              html.Small(p.description, className="text-muted")])
                     for p in rb.principles]),
        ], title=f"Principles ({len(rb.principles)})"),
        dbc.AccordionItem([
            html.Ul([html.Li(html.Code(d), style={"marginBottom": "8px"})
                     for d in rb.canonical_disclaimers]),
        ], title=f"Canonical Disclaimers ({len(rb.canonical_disclaimers)})"),
    ], start_collapsed=True, flush=False, className="mb-4")

    doc_checker = dbc.Card(dbc.CardBody([
        html.H5("Document Compliance Advisor", className="fw-bold text-primary mb-1"),
        html.P([
            "Upload any document and the advisor will review every line — titles, "
            "paragraphs, bullet points and speech elements — against the rulebook. "
            "Each violation is returned with the relevant rule, a plain-English "
            "finding, and a concrete fix."
        ], className="text-muted small mb-3"),
        dbc.Row([
            dbc.Col([
                dcc.Upload(
                    id="cp-doc-upload",
                    children=html.Div([
                        html.Span("📎 ", style={"fontSize": "1.6rem"}),
                        html.Div([
                            html.Span("Drag & drop a document or "),
                            html.A("browse to upload", style={"cursor": "pointer",
                                                               "textDecoration": "underline"}),
                        ], className="mt-1"),
                        html.Small(".docx · .txt · .md · .html",
                                   className="text-muted d-block mt-1"),
                    ], className="text-center py-4"),
                    style={
                        "border": "2px dashed #c0c8d4",
                        "borderRadius": "8px",
                        "cursor": "pointer",
                        "background": "#f8f9fa",
                        "transition": "border-color 0.15s",
                    },
                    multiple=False,
                ),
                html.Div(id="cp-doc-filename", className="text-muted small mt-2"),
            ], md=8),
            dbc.Col([
                dbc.Button(
                    [html.Span("🛡 ", style={"fontSize": "1.1rem"}), " Run Compliance Check"],
                    id="cp-doc-check-btn",
                    color="primary",
                    disabled=True,
                    size="lg",
                    className="w-100 h-100",
                    style={"minHeight": "80px"},
                ),
            ], md=4, className="d-flex align-items-center"),
        ], className="g-3"),
        html.Div(id="cp-doc-result", className="mt-4"),
        dcc.Store(id="cp-doc-job-id"),
        dcc.Store(id="cp-doc-result-store"),
        dcc.Interval(id="cp-doc-poll", interval=800, disabled=True, n_intervals=0),
    ]), className="shadow-sm mb-4")

    return html.Div([
        html.H3("Compliance", className="fw-bold mb-1"),
        html.P([
            "Loaded from ", html.Code(rb.source_path),
            " · cached at ", html.Code("data/compliance_rules.json"),
        ], className="text-muted mb-4"),
        summary_cards,
        doc_checker,
        dbc.Card(dbc.CardBody([
            html.H5("Marketing Compliance Rulebook", className="fw-bold text-primary mb-3"),
            rules_html,
        ]), className="shadow-sm mb-4"),
        dbc.Accordion([
            dbc.AccordionItem(
                log_table,
                title=f"Recent Compliance Runs ({len(log_rows)})",
            ),
            dbc.AccordionItem(
                flagged_table,
                title=f"Flagged Articles ({len(flagged_articles)})",
            ),
        ], start_collapsed=True, flush=False, className="mb-4 shadow-sm"),
    ])


# ===========================================================================
# COMPLIANCE PAGE — document checker callbacks
# ===========================================================================

def _extract_doc_text(filename: str, content_b64: str) -> tuple[str, str | None]:
    """Decode a dcc.Upload payload and return (plain_text, error_msg).

    Supports .docx, .txt, .md, .html.  Returns ("", error) for unsupported types.
    """
    import base64, io, re as _re

    # dcc.Upload sends "data:<mime>;base64,<data>"
    if "," in content_b64:
        content_b64 = content_b64.split(",", 1)[1]

    raw = base64.b64decode(content_b64)
    ext = os.path.splitext(filename.lower())[1]

    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(io.BytesIO(raw))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip()), None
        except Exception as e:
            return "", f"Could not read .docx: {e}"

    if ext in (".txt", ".md"):
        try:
            return raw.decode("utf-8", errors="replace"), None
        except Exception as e:
            return "", f"Could not decode file: {e}"

    if ext in (".html", ".htm"):
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(raw.decode("utf-8", errors="replace"), "html.parser")
            return soup.get_text(separator=" "), None
        except Exception as e:
            return "", f"Could not parse HTML: {e}"

    if ext == ".pdf":
        return "", "PDF files are not supported — please upload a .docx, .txt, .md, or .html file."

    return "", f"Unsupported file type '{ext}' — use .docx, .txt, .md, or .html."


@app.callback(
    Output("cp-doc-filename",  "children"),
    Output("cp-doc-check-btn", "disabled"),
    Input("cp-doc-upload",     "filename"),
    Input("cp-doc-upload",     "contents"),
    prevent_initial_call=True,
)
def _cp_on_upload(filename, contents):
    if not filename or not contents:
        return "", True
    _, err = _extract_doc_text(filename, contents)
    if err:
        return dbc.Alert(err, color="warning", className="mb-0 py-2 small"), True
    return html.Span(["📄 ", html.Strong(filename), " — ready"],
                     className="text-success"), False


# ---------------------------------------------------------------------------
# Shared result-rendering helpers
# ---------------------------------------------------------------------------

_SEV_COLOR = {"critical": "danger", "warning": "warning", "suggestion": "info"}
_SEV_ICON  = {"critical": "❌", "warning": "⚠️", "suggestion": "💡"}
_TYPE_ICON = {"heading": "H", "paragraph": "¶", "list_item": "•", "line": "—"}


def _cp_progress_card(job_id: str) -> html.Div:
    job = _cp_jobs.get(job_id) or {}
    stage   = job.get("stage", "parse")
    started = job.get("started_at") or time.time()
    elapsed = int(time.time() - started)

    current_idx = next((i for i, (k, _) in enumerate(_CP_STAGES) if k == stage), 0)
    rows = []
    for i, (key, label) in enumerate(_CP_STAGES):
        if i < current_idx or (key == "done" and stage == "done"):
            icon = html.Span("✅", style={"width": "22px", "display": "inline-block"})
            cls  = "text-success fw-semibold"
        elif i == current_idx:
            icon = dbc.Spinner(size="sm", color="primary",
                               spinner_style={"width": "1rem", "height": "1rem",
                                              "marginRight": "4px"})
            cls  = "text-primary fw-bold"
        else:
            icon = html.Span("○", style={"width": "22px", "display": "inline-block",
                                          "color": "#c0c8d4"})
            cls  = "text-muted"
        rows.append(html.Div([icon, html.Span(label, className=cls)], className="mb-2"))

    pct = int((current_idx / max(len(_CP_STAGES) - 1, 1)) * 100)
    if stage == "done":
        pct = 100

    return dbc.Card(dbc.CardBody([
        dbc.Row([
            dbc.Col([html.H6("Running compliance check…", className="fw-bold mb-1"),
                     html.Small(f"Elapsed: {elapsed}s", className="text-muted")], md=8),
            dbc.Col(dbc.Badge(f"{pct}%", color="primary", className="float-end fs-6"), md=4),
        ], className="mb-3"),
        dbc.Progress(value=pct, striped=True, animated=(stage != "done"),
                     className="mb-3", style={"height": "8px"}),
        html.Div(rows),
    ]), className="shadow-sm", style={"borderLeft": "4px solid #1a4f8b"})


def _cp_render_result(filename: str, result: dict, text: str) -> html.Div:
    findings   = result.get("findings") or []
    clean      = result.get("clean_count", 0)
    total      = result.get("total_count", 0)
    summary    = result.get("summary", "")
    elapsed    = result.get("elapsed_seconds", 0)
    el_lookup  = {e["index"]: e for e in compliance_parse_elements(text)}

    criticals  = sum(1 for f in findings if f.get("severity") == "critical")
    warnings_n = sum(1 for f in findings if f.get("severity") == "warning")
    suggests   = sum(1 for f in findings if f.get("severity") == "suggestion")

    overall_color = "success" if not findings else ("danger" if criticals else "warning")
    overall_icon  = "✅" if not findings else ("❌" if criticals else "⚠️")

    stat_badges = []
    if criticals:
        stat_badges.append(dbc.Badge(f"{criticals} critical",    color="danger",    className="me-2"))
    if warnings_n:
        stat_badges.append(dbc.Badge(f"{warnings_n} warning(s)", color="warning",   className="me-2"))
    if suggests:
        stat_badges.append(dbc.Badge(f"{suggests} suggestion(s)", color="info",     className="me-2"))
    if not findings:
        stat_badges.append(dbc.Badge("All clear", color="success"))

    banner = dbc.Alert([
        dbc.Row([
            dbc.Col([
                html.Span(overall_icon + " ", style={"fontSize": "1.3rem"}),
                html.Strong(filename),
                html.Span(f"  ·  {total} elements reviewed  ·  {clean} clean  ·  {elapsed}s",
                           className="text-muted ms-2 small"),
                html.Div(stat_badges, className="mt-2"),
                html.P(summary, className="mb-0 mt-2 small"),
            ], md=9),
            dbc.Col([
                dbc.Button(
                    "⬇ Download Report (.docx)",
                    id="cp-doc-download-btn",
                    color="light",
                    size="sm",
                    className="w-100",
                    style={"fontWeight": "600"},
                ),
                html.Div(id="cp-doc-download-link", className="mt-2 text-center"),
            ], md=3, className="d-flex flex-column justify-content-center"),
        ]),
    ], color=overall_color, className="mb-3")

    if not findings:
        return html.Div([banner,
                         html.Div(id="cp-doc-download-btn",  style={"display": "none"}),
                         html.Div(id="cp-doc-download-link", style={"display": "none"})])

    def _finding_row(f: dict) -> dbc.Card:
        sev   = f.get("severity", "suggestion")
        idx   = f.get("index", 0)
        el    = el_lookup.get(idx, {})
        etype = el.get("type", "line")
        etext = el.get("text", "")
        return dbc.Card(dbc.CardBody([
            dbc.Row([
                dbc.Col([
                    dbc.Badge(_SEV_ICON.get(sev, "•") + " " + sev.upper(),
                              color=_SEV_COLOR.get(sev, "secondary"),
                              className="me-2 mb-1", style={"fontSize": "0.72rem"}),
                    dbc.Badge(_TYPE_ICON.get(etype, "—") + " " + etype,
                              color="light", text_color="dark",
                              style={"fontSize": "0.72rem"}),
                ], md=3, className="d-flex align-items-start flex-wrap gap-1 pt-1"),
                dbc.Col([
                    html.Div(f'"{etext[:160]}{"…" if len(etext) > 160 else ""}"',
                             className="text-muted small fst-italic mb-1"),
                    html.Div([html.Strong("Finding: ", className="small"),
                              html.Span(f.get("finding", ""), className="small")]),
                    html.Div([
                        dbc.Badge(f"§{f['section']}", color="warning", className="me-1 mt-1",
                                  style={"fontSize": "0.68rem"}) if f.get("section") else html.Span(),
                        dbc.Badge(f.get("rule", ""), color="light", text_color="dark",
                                  style={"fontSize": "0.68rem"}),
                    ], className="mt-1"),
                ], md=5),
                dbc.Col([
                    html.Div("Fix:", className="text-muted small fw-semibold mb-1"),
                    html.Div(f.get("solution", ""), className="small",
                             style={"lineHeight": "1.4"}),
                ], md=4, style={"borderLeft": "3px solid #e9ecef", "paddingLeft": "12px"}),
            ], className="g-2 align-items-start"),
        ]), className="mb-2 shadow-sm",
           style={"borderLeft": f"4px solid var(--bs-{_SEV_COLOR.get(sev, 'secondary')})"})

    return html.Div([banner, html.Div([_finding_row(f) for f in findings])])


# ---------------------------------------------------------------------------
# Report .docx generator
# ---------------------------------------------------------------------------

def _build_compliance_report_docx(filename: str, result: dict, text: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    findings  = result.get("findings") or []
    clean     = result.get("clean_count", 0)
    total     = result.get("total_count", 0)
    summary   = result.get("summary", "")
    elapsed   = result.get("elapsed_seconds", 0)
    el_lookup = {e["index"]: e for e in compliance_parse_elements(text)}

    criticals  = sum(1 for f in findings if f.get("severity") == "critical")
    warnings_n = sum(1 for f in findings if f.get("severity") == "warning")
    suggests   = sum(1 for f in findings if f.get("severity") == "suggestion")

    _SEV_RGB = {
        "critical":   RGBColor(0xC0, 0x39, 0x2B),
        "warning":    RGBColor(0xE6, 0x7E, 0x22),
        "suggestion": RGBColor(0x29, 0x80, 0xB9),
    }
    _GREY     = RGBColor(0x7F, 0x8C, 0x8D)
    _DARK     = RGBColor(0x2C, 0x3E, 0x50)
    _WARREN   = RGBColor(0x1A, 0x4F, 0x8B)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ---- Title block -------------------------------------------------------
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = t.add_run("Compliance Check Report")
    run.font.size  = Pt(22)
    run.font.bold  = True
    run.font.color.rgb = _WARREN

    sub = doc.add_paragraph()
    r = sub.add_run(filename)
    r.font.size  = Pt(12)
    r.font.color.rgb = _GREY

    meta = doc.add_paragraph()
    r = meta.add_run(
        f"Generated: {datetime.utcnow().strftime('%d %B %Y, %H:%M UTC')}  ·  "
        f"Checked in {elapsed}s"
    )
    r.font.size  = Pt(9)
    r.font.color.rgb = _GREY

    doc.add_paragraph()

    # ---- Summary box (table-as-box) ----------------------------------------
    doc.add_heading("Summary", level=2)

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    cells = tbl.rows[0].cells
    for cell, label, val, rgb in [
        (cells[0], "Total Elements", str(total),     _DARK),
        (cells[1], "Clean",          str(clean),      RGBColor(0x27, 0xAE, 0x60)),
        (cells[2], "Issues Found",   str(len(findings)), _SEV_RGB.get("warning", _DARK)),
        (cells[3], "Critical",       str(criticals),  _SEV_RGB.get("critical", _DARK)),
    ]:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(val + "\n")
        r1.font.size  = Pt(18)
        r1.font.bold  = True
        r1.font.color.rgb = rgb
        r2 = p.add_run(label)
        r2.font.size  = Pt(8)
        r2.font.color.rgb = _GREY

    doc.add_paragraph()

    sev_line = doc.add_paragraph()
    for label, count, rgb in [
        ("Critical", criticals,  _SEV_RGB["critical"]),
        ("Warning",  warnings_n, _SEV_RGB["warning"]),
        ("Suggestion", suggests, _SEV_RGB["suggestion"]),
    ]:
        r = sev_line.add_run(f"  {label}: {count}  ")
        r.font.bold  = True
        r.font.color.rgb = rgb

    p = doc.add_paragraph(summary)
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = _GREY

    doc.add_paragraph()

    # ---- Findings ---------------------------------------------------------
    if findings:
        doc.add_heading("Findings", level=2)

        for i, f in enumerate(findings, 1):
            sev   = f.get("severity", "suggestion")
            idx   = f.get("index", 0)
            el    = el_lookup.get(idx, {})
            etype = el.get("type", "line")
            etext = el.get("text", "")
            rgb   = _SEV_RGB.get(sev, _DARK)

            # Finding heading
            hdr = doc.add_paragraph()
            rn = hdr.add_run(f"{i}.  [{sev.upper()}]  {f.get('rule', '')}  ")
            rn.font.bold  = True
            rn.font.color.rgb = rgb
            if f.get("section"):
                rs = hdr.add_run(f"§{f['section']}")
                rs.font.size  = Pt(9)
                rs.font.color.rgb = _GREY

            # Element type + original text
            orig = doc.add_paragraph()
            rt = orig.add_run(f"{etype.upper()}:  ")
            rt.font.bold  = True
            rt.font.size  = Pt(9)
            rt.font.color.rgb = _GREY
            ro = orig.add_run(f'"{etext[:300]}"')
            ro.font.italic = True
            ro.font.size   = Pt(9)
            ro.font.color.rgb = _GREY

            # Finding
            fp = doc.add_paragraph()
            fp.add_run("Finding:  ").font.bold = True
            fp.add_run(f.get("finding", ""))

            # Solution (indented, coloured)
            sp = doc.add_paragraph()
            sp.paragraph_format.left_indent = Inches(0.3)
            sr = sp.add_run("Fix:  ")
            sr.font.bold  = True
            sr.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
            sp.add_run(f.get("solution", ""))

            doc.add_paragraph()
    else:
        doc.add_paragraph("No compliance issues were found. All elements are fully compliant.")

    # ---- Footer -----------------------------------------------------------
    doc.add_page_break()
    ft = doc.add_paragraph()
    ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ft.add_run("Generated by Warren Compliance Advisor  ·  meetwarren.co.uk")
    r.font.size  = Pt(8)
    r.font.color.rgb = _GREY

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Kickoff callback — starts background job, returns progress card
# ---------------------------------------------------------------------------

@app.callback(
    Output("cp-doc-result",  "children",  allow_duplicate=True),
    Output("cp-doc-job-id",  "data"),
    Output("cp-doc-poll",    "disabled",  allow_duplicate=True),
    Output("cp-doc-check-btn", "disabled", allow_duplicate=True),
    Input("cp-doc-check-btn", "n_clicks"),
    State("cp-doc-upload",   "contents"),
    State("cp-doc-upload",   "filename"),
    prevent_initial_call=True,
)
def _cp_kickoff(_n, contents, filename):
    if not contents or not filename:
        return dbc.Alert("No file uploaded.", color="warning"), no_update, True, False

    text, err = _extract_doc_text(filename, contents)
    if err:
        return dbc.Alert(err, color="danger"), no_update, True, False
    if not text.strip():
        return dbc.Alert("File appears to be empty.", color="warning"), no_update, True, False

    job_id = uuid.uuid4().hex[:12]
    with _cp_jobs_lock:
        _cp_jobs[job_id] = {
            "stage": "parse", "started_at": time.time(),
            "error": None, "result": None, "text": text, "filename": filename,
        }

    def _worker():
        try:
            with _cp_jobs_lock:
                _cp_jobs[job_id]["stage"] = "parse"
            elements = compliance_parse_elements(text)

            with _cp_jobs_lock:
                _cp_jobs[job_id]["stage"] = "hard_rules"
            rulebook = load_rulebook()

            with _cp_jobs_lock:
                _cp_jobs[job_id]["stage"] = "llm"
            client = build_anthropic_client(cfg)
            result = advise_document(
                elements, rulebook=rulebook, client=client,
                model=cfg.compliance_model or cfg.anthropic_model,
            )

            with _cp_jobs_lock:
                _cp_jobs[job_id]["stage"]  = "done"
                _cp_jobs[job_id]["result"] = result
        except Exception as e:
            with _cp_jobs_lock:
                _cp_jobs[job_id]["stage"] = "done"
                _cp_jobs[job_id]["error"] = f"{type(e).__name__}: {e}"

    threading.Thread(target=_worker, daemon=True).start()
    return _cp_progress_card(job_id), job_id, False, True


# ---------------------------------------------------------------------------
# Poll callback — refreshes progress or renders final result
# ---------------------------------------------------------------------------

@app.callback(
    Output("cp-doc-result",       "children",  allow_duplicate=True),
    Output("cp-doc-poll",         "disabled",  allow_duplicate=True),
    Output("cp-doc-check-btn",    "disabled",  allow_duplicate=True),
    Output("cp-doc-result-store", "data"),
    Input("cp-doc-poll",   "n_intervals"),
    State("cp-doc-job-id", "data"),
    prevent_initial_call=True,
)
def _cp_poll(_n, job_id):
    if not job_id or job_id not in _cp_jobs:
        return no_update, True, False, no_update

    with _cp_jobs_lock:
        job = dict(_cp_jobs[job_id])

    if job.get("error"):
        _cp_jobs.pop(job_id, None)
        return (dbc.Alert([html.Strong("Compliance check failed: "), job["error"]],
                          color="danger"),
                True, False, no_update)

    if job.get("stage") == "done" and job.get("result") is not None:
        result   = job["result"]
        text     = job["text"]
        filename = job["filename"]
        _cp_jobs.pop(job_id, None)
        rendered = _cp_render_result(filename, result, text)
        store    = {"result": result, "text": text, "filename": filename}
        return rendered, True, False, store

    return _cp_progress_card(job_id), False, True, no_update


# ---------------------------------------------------------------------------
# Download callback — builds .docx from stored result and serves it
# ---------------------------------------------------------------------------

@app.callback(
    Output("cp-doc-download-link", "children"),
    Input("cp-doc-download-btn",   "n_clicks"),
    State("cp-doc-result-store",   "data"),
    prevent_initial_call=True,
)
def _cp_download_report(_n, store):
    if not store:
        return dbc.Alert("No result to download yet.", color="warning", className="small py-1")

    result   = store.get("result", {})
    text     = store.get("text", "")
    filename = store.get("filename", "document")

    try:
        docx_bytes = _build_compliance_report_docx(filename, result, text)
    except Exception as e:
        return dbc.Alert(f"Report generation failed: {e}", color="danger", className="small py-1")

    stem = os.path.splitext(os.path.basename(filename))[0]
    out_name = f"compliance-report-{stem}-{datetime.utcnow().strftime('%Y%m%d-%H%M')}.docx"
    out_path = os.path.join(cfg.output_dir, out_name)
    os.makedirs(cfg.output_dir, exist_ok=True)
    with open(out_path, "wb") as f:
        f.write(docx_bytes)

    return dbc.Button(
        f"⬇ {out_name}",
        href=f"/downloads/{out_name}",
        external_link=True,
        download=out_name,
        color="success",
        size="sm",
        className="w-100",
    )


# ===========================================================================
# ANSWER MACHINE PAGE — drafts brand-aligned replies to incoming comments/DMs
# ===========================================================================

# Module-level KB cache so we don't re-parse the docx/xlsx on every request.
# `load_knowledge_base()` itself checks source-file mtimes and rebuilds when
# stale, so this just avoids the repeated JSON read.
import threading as _am_threading
_am_kb_lock = _am_threading.Lock()
_am_kb_cache: dict = {"kb": None}


def _am_get_kb(force: bool = False):
    with _am_kb_lock:
        if force or _am_kb_cache["kb"] is None:
            _am_kb_cache["kb"] = load_knowledge_base(force_rebuild=force)
        return _am_kb_cache["kb"]


# Pre-canned message templates — recurring patterns lifted directly from the
# Meta corpus. Click → pre-fills the paste box (and optionally platform/kind).
# Useful for QA, training new team members, and reproducing edge cases.
_AM_TEMPLATES: list[dict] = [
    {
        "label":   "Just use ChatGPT",
        "message": "Just use the latest version of ChatGPT in thinking mode",
        "platform": "Instagram", "kind": "comment",
        "tooltip": "Generic 'why bother with Warren' challenge — most common comment.",
    },
    {
        "label":   "Just speak to an IFA",
        "message": "Do not use - far better speak to a local IFA and get regulated advice - not vague and often incorrect information",
        "platform": "Facebook", "kind": "comment",
        "tooltip": "Anti-AI / pro-IFA pushback. Standard reply uses the dentist/toothpaste analogy.",
    },
    {
        "label":   "It's just a spreadsheet",
        "message": "\"Not Advice\" - so basically an Excel Spreadsheet to add 1 and 1 to make 2????",
        "platform": "Facebook", "kind": "comment",
        "tooltip": "Dismissive 'this is just Excel' jab.",
    },
    {
        "label":   "Not FCA approved",
        "message": "Good idea, but not FCA approved, it's AI with extra steps. What is your security? Overall liability for your financial advice? How do I get in touch with a person?",
        "platform": "Instagram", "kind": "comment",
        "tooltip": "Constructive scrutiny — security + advice scope + escalation path.",
    },
    {
        "label":   "Disclaimer contradiction",
        "message": "You claim this is \"not financial advice,\" yet you go on to offer what is clearly financial guidance. as if adding a caveat somehow makes it acceptable. People should be very wary of products like this as there is no protection if things go wrong.",
        "platform": "Instagram", "kind": "comment",
        "tooltip": "Legal/regulatory challenge to the 'not advice' line.",
    },
    {
        "label":   "Pricing",
        "message": "cost?",
        "platform": "Instagram", "kind": "comment",
        "tooltip": "Quick pricing question — short reply with the canonical pricing block.",
    },
    {
        "label":   "Is it free?",
        "message": "Is it free",
        "platform": "Instagram", "kind": "comment",
        "tooltip": "Free-tier question.",
    },
    {
        "label":   "Android app?",
        "message": "Will there be an android app?",
        "platform": "Instagram", "kind": "dm",
        "tooltip": "Roadmap question — DM, conversational reply.",
    },
    {
        "label":   "Login problem",
        "message": "I get login failed error when trying to sign up",
        "platform": "Instagram", "kind": "dm",
        "tooltip": "Support / troubleshooting DM.",
    },
    {
        "label":   "Hostile / troll",
        "message": "Looking forward to you being bankrupt after fooling thousands of idiots to trust your useless robot to handle their finances.",
        "platform": "Instagram", "kind": "comment",
        "tooltip": "Hostile commentary — calm, factual, doesn't escalate.",
    },
]


def _am_template_buttons() -> html.Div:
    """Pre-canned message templates as a wrap-friendly button row."""
    btns = []
    for i, tpl in enumerate(_AM_TEMPLATES):
        bid = {"type": "am-template", "index": i}
        btns.append(html.Span([
            dbc.Button(tpl["label"], id=bid, color="light", size="sm",
                       outline=True, n_clicks=0,
                       className="me-2 mb-2",
                       style={"fontSize": "12px", "fontWeight": "600"}),
            dbc.Tooltip(tpl["tooltip"], target=bid, placement="top"),
        ]))
    return html.Div([
        html.Div("Quick templates · click to load", className="section-eyebrow"),
        html.Div(btns, style={"display": "flex", "flexWrap": "wrap"}),
    ], className="mb-3")


# ---------------------------------------------------------------------------
# KB browser tab (I) — searchable view of FAQs + past replies
# ---------------------------------------------------------------------------

_AM_SENTIMENT_COLORS = {
    "positive":     "success",
    "question":     "info",
    "constructive": "primary",
    "-ive":         "danger",
    "misc":         "secondary",
    "approved":     "success",
}


def _am_kb_browser_layout(kb) -> html.Div:
    """Render the search/filter shell.

    The two list panes are PRE-POPULATED with the full KB at page load (no
    callback needed for initial render — avoids dash 4 initial-fire
    flakiness). The Search button + dropdown changes update them via the
    `_am_kb_filter` callback.
    """
    sections   = sorted({f.section for f in kb.faqs if f.section})
    platforms  = sorted({c.platform for c in kb.comment_examples if c.platform})
    sentiments = sorted({c.sentiment for c in kb.comment_examples if c.sentiment})

    return html.Div([
        # --- Filter strip ---------------------------------------------------
        dbc.Card(dbc.CardBody([
            html.Div("Search the knowledge base", className="section-eyebrow"),
            dbc.Row([
                dbc.Col([
                    dbc.Label("Search text"),
                    dbc.Input(id="am-kb-search", type="text",
                              placeholder="e.g. ISA, advice, ChatGPT, pension…",
                              debounce=True, n_submit=0),
                    html.Small("Tip: press Enter to search.",
                               className="text-muted"),
                ], md=5),
                dbc.Col([
                    dbc.Label("FAQ section"),
                    dcc.Dropdown(id="am-kb-section",
                                 options=[{"label": s, "value": s} for s in sections],
                                 placeholder="All sections", multi=True, clearable=True),
                ], md=3),
                dbc.Col([
                    dbc.Label("Platform"),
                    dcc.Dropdown(id="am-kb-platform",
                                 options=[{"label": p, "value": p} for p in platforms],
                                 placeholder="All platforms", multi=True, clearable=True),
                ], md=2),
                dbc.Col([
                    dbc.Label("Sentiment"),
                    dcc.Dropdown(id="am-kb-sentiment",
                                 options=[{"label": s, "value": s} for s in sentiments],
                                 placeholder="All", multi=True, clearable=True),
                ], md=2),
            ], className="g-3"),
            html.Div([
                dbc.Button("🔍  Search", id="am-kb-search-btn",
                           color="primary", size="sm", className="me-2"),
                dbc.Button("✕  Clear", id="am-kb-reset-btn",
                           color="light", size="sm", outline=True),
            ], className="mt-3"),
        ]), className="mb-3"),

        # Toast slot for delete feedback — sits OUTSIDE the replies pane so
        # the message survives when the pane re-renders post-delete.
        html.Div(id="am-kb-delete-status", className="mb-2"),

        # --- Two stacked sections; pre-populated with the full corpus -------
        html.Div(_am_render_faqs(kb.faqs),
                 id="am-kb-faqs-pane",     className="mb-3"),
        html.Div(_am_render_replies(kb.comment_examples),
                 id="am-kb-replies-pane"),
    ])


def _am_render_faqs(faqs: list) -> html.Div:
    if not faqs:
        return dbc.Card(dbc.CardBody([
            html.Div("FAQs", className="section-eyebrow"),
            html.Em("No FAQs match your filters.", className="text-muted"),
        ]))
    items = []
    for f in faqs:
        title = html.Div([
            html.Strong(f.question),
            html.Span(f"  ·  {f.section}" if f.section else "",
                      className="text-muted small ms-2"),
        ])
        body = html.Div([
            (html.Div([
                html.Div("Short answer", className="section-eyebrow"),
                html.Div(f.answer_short, style={"whiteSpace": "pre-wrap",
                                                 "marginBottom": "12px"}),
            ]) if f.answer_short else html.Div()),
            (html.Div([
                html.Div("Long answer", className="section-eyebrow"),
                html.Div(f.answer_long, style={"whiteSpace": "pre-wrap"}),
            ]) if f.answer_long else html.Div()),
        ])
        items.append(dbc.AccordionItem(body, title=title))

    return dbc.Card(dbc.CardBody([
        html.Div([
            html.Div(f"FAQs · {len(faqs)} match{'es' if len(faqs) != 1 else ''}",
                     className="section-eyebrow d-inline"),
        ], className="mb-2"),
        dbc.Accordion(items, start_collapsed=True, flush=True,
                      always_open=False),
    ]))


def _am_render_replies(replies: list) -> html.Div:
    if not replies:
        return dbc.Card(dbc.CardBody([
            html.Div("Past replies", className="section-eyebrow"),
            html.Em("No past replies match your filters.", className="text-muted"),
        ]))
    items = []
    for c in replies:
        sentiment_color = _AM_SENTIMENT_COLORS.get((c.sentiment or "").strip(),
                                                    "secondary")
        title = html.Div([
            dbc.Badge(c.platform or "?", color="primary",
                      className="me-1", style={"fontSize": "0.65rem"}),
            dbc.Badge((c.sentiment or "?").upper(), color=sentiment_color,
                      className="me-1", style={"fontSize": "0.65rem"}),
            dbc.Badge("DM" if c.is_dm else "COMMENT", color="info",
                      className="me-2", style={"fontSize": "0.65rem"}),
            html.Span(c.comment[:120] + ("…" if len(c.comment) > 120 else ""),
                      style={"fontSize": "13px"}),
            (html.Span(f"  ·  {c.date}", className="text-muted small ms-1")
             if c.date else html.Span()),
        ])
        # Delete control: ConfirmDialogProvider shows a browser-native confirm
        # so a stray click can't wipe a row. Identifier is the content hash,
        # not an array index — so deleting one row doesn't mis-target the next.
        delete_ctl = (
            dcc.ConfirmDialogProvider(
                children=dbc.Button(
                    "🗑  Delete this exemplar",
                    color="danger", size="sm", outline=True,
                ),
                id={"type": "am-kb-delete", "row_id": c.row_id or ""},
                message=("Delete this past reply from the knowledge base?\n\n"
                         "This rewrites meta_comments.xlsx and cannot be undone."),
            ) if c.row_id else html.Div()
        )
        body = html.Div([
            html.Div("Incoming", className="section-eyebrow"),
            html.Div(c.comment, style={"whiteSpace": "pre-wrap",
                                        "marginBottom": "12px"}),
            html.Div("Our reply", className="section-eyebrow"),
            html.Div(c.response, style={"whiteSpace": "pre-wrap",
                                         "fontStyle": "italic",
                                         "background": "var(--warren-soft)",
                                         "padding": "10px 14px",
                                         "borderRadius": "var(--radius-sm)",
                                         "border": "1px solid var(--warren-border)"}),
            html.Div([
                html.Span(f"— @{c.account}" if c.account else "",
                          className="text-muted small me-3"),
                delete_ctl,
            ], className="d-flex justify-content-between align-items-center mt-3"),
        ])
        items.append(dbc.AccordionItem(body, title=title))

    return dbc.Card(dbc.CardBody([
        html.Div(f"Past replies · {len(replies)} match{'es' if len(replies) != 1 else ''}",
                 className="section-eyebrow mb-2"),
        dbc.Accordion(items, start_collapsed=True, flush=True),
    ]))


def _answer_machine_page():
    try:
        kb = _am_get_kb()
        kb_status = dbc.Badge(
            f"KB loaded · {len(kb.faqs)} FAQs · {len(kb.comment_examples)} past replies",
            color="success", className="me-2",
        )
        kb_error = None
    except Exception as e:
        kb = None
        kb_status = dbc.Badge("KB load failed", color="danger", className="me-2")
        kb_error  = str(e)

    # ---- Draft tab ---------------------------------------------------------
    left_pane = html.Div([
        dbc.Card(dbc.CardBody([
            html.Div("Step 1 · Paste the message", className="section-eyebrow"),
            dbc.Row([
                dbc.Col([
                    dbc.Label("Platform"),
                    dcc.Dropdown(id="am-platform",
                                 options=[
                                     {"label": "Instagram",  "value": "Instagram"},
                                     {"label": "Facebook",   "value": "Facebook"},
                                     {"label": "TikTok",     "value": "TikTok"},
                                     {"label": "LinkedIn",   "value": "LinkedIn"},
                                     {"label": "X / Twitter", "value": "X"},
                                     {"label": "Other",      "value": "Other"},
                                 ],
                                 value="Instagram", clearable=False),
                ], md=5),
                dbc.Col([
                    dbc.Label("Message type"),
                    dbc.RadioItems(
                        id="am-kind",
                        options=[
                            {"label": "Public comment", "value": "comment"},
                            {"label": "Direct message", "value": "dm"},
                        ],
                        value="comment", inline=True,
                        inputClassName="me-1", labelClassName="me-3",
                    ),
                ], md=7),
            ], className="g-3"),

            html.Hr(style={"margin": "16px 0"}),
            _am_template_buttons(),

            dbc.Label("Incoming message", className="mt-2"),
            dbc.Textarea(
                id="am-message",
                placeholder="Paste the comment or DM here, or pick a template above…",
                style={"minHeight": "180px", "fontSize": "14px",
                       "fontFamily": "var(--font-body)"},
            ),
            html.Div([
                dbc.Button("✨  Draft Warren's reply",
                           id="am-generate-btn", color="primary",
                           size="lg", className="mt-3"),
                dbc.Button("Refresh KB",
                           id="am-refresh-kb-btn", color="link",
                           size="sm", className="mt-3 ms-2"),
            ]),
        ]), className="mb-3"),
    ], className="create-pane-left")

    right_pane = html.Div([
        html.Div([kb_status,
                  html.Span("Powered by Brand Narrative + FAQs + Meta exemplar replies",
                            className="text-muted small")],
                 className="mb-2"),
        html.Div(id="am-output",
                 children=html.Div("Draft a reply to see it here.",
                                   className="text-muted small p-4 text-center")),
    ], className="create-pane-right")

    draft_tab_body = html.Div([
        html.Div([left_pane, right_pane], className="create-shell"),
    ])

    # ---- Browse KB tab -----------------------------------------------------
    browse_tab_body = _am_kb_browser_layout(kb) if kb else html.Div()

    return html.Div([
        _page_header("Answer Machine",
                     "Paste a comment or DM. Get a Warren-tone reply grounded "
                     "in the FAQs, brand narrative, and 30+ past replies."),
        (dbc.Alert(f"Knowledge base error: {kb_error}", color="danger")
         if kb_error else html.Div()),
        # Persists the in-flight draft (incoming message + metadata) across
        # callbacks so the Approve button can write the right exemplar.
        dcc.Store(id="am-current-draft", data=None),
        dbc.Tabs([
            dbc.Tab(draft_tab_body,  label="✍ Draft a reply",  tab_id="am-tab-draft"),
            dbc.Tab(browse_tab_body, label="🗂 Browse KB",     tab_id="am-tab-kb"),
        ], id="am-tabs", active_tab="am-tab-draft", className="mb-3"),
    ])


def _am_render_reply(out: dict, platform: str, is_dm: bool) -> html.Div:
    """Render the result of draft_reply into a card stack."""
    if out.get("error"):
        return dbc.Alert(
            [html.Strong("Drafting failed: "), out["error"]],
            color="danger",
        )
    reply = out.get("reply") or "(empty reply — try again)"
    matched_faqs     = out.get("matched_faqs", [])
    matched_examples = out.get("matched_examples", [])

    # Sources accordion
    src_items = []
    if matched_faqs:
        src_items.append(dbc.AccordionItem([
            html.Ol([
                html.Li([
                    html.Strong(f["question"]), html.Br(),
                    html.Small(f.get("answer_short") or f.get("answer_long", ""),
                               className="text-muted"),
                ], style={"marginBottom": "10px"})
                for f in matched_faqs
            ]),
        ], title=f"Top {len(matched_faqs)} FAQ matches"))
    if matched_examples:
        src_items.append(dbc.AccordionItem([
            html.Ol([
                html.Li([
                    dbc.Badge(c.get("platform") or "?", color="secondary",
                              className="me-1", style={"fontSize": "0.65rem"}),
                    dbc.Badge(c.get("sentiment") or "?",
                              color={"positive": "success", "question": "info",
                                     "constructive": "primary",
                                     "-ive": "danger", "misc": "secondary"}.get(
                                          c.get("sentiment"), "secondary"),
                              className="me-1", style={"fontSize": "0.65rem"}),
                    html.Span(("DM " if c.get("is_dm") else "Comment "),
                              className="text-muted small me-1"),
                    html.Br(),
                    html.Small("Incoming: ", className="text-muted"),
                    c["comment"][:240],
                    html.Br(),
                    html.Small("Reply: ", className="text-muted"),
                    html.Span(c["response"][:280],
                              style={"fontStyle": "italic"}),
                ], style={"marginBottom": "12px", "fontSize": "13px"})
                for c in matched_examples
            ]),
        ], title=f"Top {len(matched_examples)} past-reply matches"))

    cache_label = ("⚡ cache hit" if out.get("cache_hit")
                   else "(no cache hit)" if out.get("cache_hit") is False
                   else "")

    return html.Div([
        # Reply card with editable textarea + copy + approve
        dbc.Card(dbc.CardBody([
            dbc.Row([
                dbc.Col([
                    html.Div("Drafted reply (editable)", className="section-eyebrow"),
                    html.Div([
                        dbc.Badge(platform, color="primary", className="me-1"),
                        dbc.Badge("DM" if is_dm else "Comment", color="info",
                                  className="me-1"),
                        html.Small(f"{out.get('elapsed_seconds', '?')}s · "
                                   f"{out.get('model','?').split('-')[1] if '-' in out.get('model','') else out.get('model','')}"
                                   f" · {cache_label}",
                                   className="text-muted"),
                    ], className="mb-2"),
                ], md=9),
                dbc.Col([
                    dcc.Clipboard(target_id="am-reply-text",
                                   title="Copy to clipboard",
                                   style={"fontSize": "1.4rem",
                                          "cursor": "pointer",
                                          "float": "right"}),
                ], md=3),
            ]),
            # Editable textarea so the team can polish before approving.
            # The id is also the clipboard target — dcc.Clipboard reads .value
            # for textarea elements automatically.
            dbc.Textarea(
                id="am-reply-text",
                value=reply,
                style={"fontFamily": "var(--font-body)",
                       "fontSize": "14px", "lineHeight": "1.55",
                       "background": "var(--warren-soft)",
                       "padding": "16px", "borderRadius": "var(--radius)",
                       "border": "1px solid var(--warren-border)",
                       "color": "var(--warren-ink)",
                       "minHeight": "260px", "resize": "vertical"},
            ),

            # Approve & add to KB row
            dbc.Row([
                dbc.Col([
                    dbc.Label("Tag this exemplar as", className="small text-muted mb-1"),
                    dcc.Dropdown(
                        id="am-approve-sentiment",
                        options=[
                            {"label": "Approved (default)", "value": "approved"},
                            {"label": "Question",           "value": "question"},
                            {"label": "Constructive",       "value": "constructive"},
                            {"label": "Positive",           "value": "positive"},
                            {"label": "Negative / hostile", "value": "-ive"},
                            {"label": "Misc",               "value": "misc"},
                        ],
                        value="approved", clearable=False,
                        style={"fontSize": "13px"},
                    ),
                ], md=5),
                dbc.Col([
                    dbc.Button("✓  Approve & add to KB",
                               id="am-approve-btn", color="success",
                               className="mt-3 w-100"),
                    dbc.Tooltip(
                        "Saves the (incoming, edited reply) pair to the Meta Comments "
                        "spreadsheet so future drafts can learn from it. "
                        "The reply text used is whatever is currently in the editor above.",
                        target="am-approve-btn", placement="top",
                    ),
                ], md=4),
                dbc.Col([
                    html.Div(id="am-approve-status", className="mt-3"),
                ], md=3),
            ], className="mt-3 g-2"),
        ]), className="mb-3 shadow-sm",
            style={"borderLeft": "4px solid var(--bs-primary)"}),

        # Sources accordion
        (dbc.Card(dbc.CardBody([
            html.Div("Sources used", className="section-eyebrow"),
            dbc.Accordion(src_items, start_collapsed=True, flush=True),
        ]), className="shadow-sm") if src_items else html.Div()),
    ])


@app.callback(
    Output("am-output", "children"),
    Output("am-generate-btn", "disabled"),
    Output("am-current-draft", "data"),
    Input("am-generate-btn", "n_clicks"),
    State("am-message",  "value"),
    State("am-platform", "value"),
    State("am-kind",     "value"),
    prevent_initial_call=True,
)
def _am_draft(_n, message, platform, kind):
    if not message or not message.strip():
        return dbc.Alert("Paste a message first.", color="warning"), False, no_update
    try:
        kb = _am_get_kb()
        client = build_anthropic_client(cfg)
    except Exception as e:
        return dbc.Alert(f"Setup failed: {e}", color="danger"), False, no_update
    is_dm = (kind == "dm")
    out = draft_reply(message, client=client, model=cfg.anthropic_model,
                      kb=kb, platform_hint=platform, is_dm=is_dm)
    # Stash the inputs so the Approve button knows what to write.
    draft_state = {
        "message":  message.strip(),
        "platform": platform or "",
        "is_dm":    is_dm,
    }
    return _am_render_reply(out, platform=platform, is_dm=is_dm), False, draft_state


@app.callback(
    Output("am-output", "children", allow_duplicate=True),
    Input("am-refresh-kb-btn", "n_clicks"),
    prevent_initial_call=True,
)
def _am_refresh(_n):
    try:
        kb = _am_get_kb(force=True)
        return dbc.Alert(
            f"Knowledge base rebuilt: {len(kb.faqs)} FAQs, "
            f"{len(kb.comment_examples)} past replies, "
            f"{len(kb.brand_voice_principles)} voice principles.",
            color="success",
        )
    except Exception as e:
        return dbc.Alert(f"Rebuild failed: {e}", color="danger")


@app.callback(
    Output("am-approve-status",   "children"),
    Output("am-approve-btn",      "disabled"),
    Input("am-approve-btn",       "n_clicks"),
    State("am-reply-text",        "value"),
    State("am-approve-sentiment", "value"),
    State("am-current-draft",     "data"),
    prevent_initial_call=True,
)
def _am_approve(_n, current_reply, sentiment, draft_state):
    """Write the (incoming, edited reply) pair as a new exemplar to the
    Meta Comments spreadsheet. The KB is rebuilt automatically by the writer
    so the next draft sees the new example.
    """
    if not draft_state:
        return dbc.Alert("Generate a draft first.", color="warning",
                          className="mb-0 py-2 small"), False
    if not current_reply or not current_reply.strip():
        return dbc.Alert("Reply is empty — write something to save.",
                          color="warning", className="mb-0 py-2 small"), False

    result = append_exemplar(
        comment=draft_state.get("message", ""),
        response=current_reply,
        platform=draft_state.get("platform", ""),
        sentiment=sentiment or "approved",
        is_dm=bool(draft_state.get("is_dm")),
    )
    if not result.get("ok"):
        return dbc.Alert(f"⚠ {result.get('error', 'unknown error')}",
                          color="danger", className="mb-0 py-2 small"), False

    # Update the module-level cached KB so the very next draft sees the change.
    with _am_kb_lock:
        _am_kb_cache["kb"] = result["kb"]

    n = len(result["kb"].comment_examples)
    return dbc.Alert(
        [
            html.Strong("✓ Added to KB"),
            html.Br(),
            html.Small(f"Now {n} exemplar replies. The next draft will use this one.",
                       className="text-muted"),
        ],
        color="success", className="mb-0 py-2 small",
    ), True   # disable the button so it can't be clicked twice on the same draft


# --- Template buttons (H) ---------------------------------------------------

@app.callback(
    Output("am-message",  "value"),
    Output("am-platform", "value"),
    Output("am-kind",     "value"),
    Input({"type": "am-template", "index": dash.ALL}, "n_clicks"),
    prevent_initial_call=True,
)
def _am_template_click(_clicks):
    """Pattern-matching callback: any template button → pre-fill the input.

    NOTE on the lack of allow_duplicate: dash 4 validates allow_duplicate=True
    strictly and silently drops updates when no other callback writes the
    same output. Nothing else writes am-message/platform/kind, so we don't
    need (or want) the flag here.
    """
    triggered = ctx.triggered_id
    if not triggered or not _clicks or not any(c for c in _clicks if c):
        return no_update, no_update, no_update
    idx = triggered.get("index") if isinstance(triggered, dict) else None
    if idx is None or idx >= len(_AM_TEMPLATES):
        return no_update, no_update, no_update
    tpl = _AM_TEMPLATES[idx]
    return (
        tpl["message"],
        tpl.get("platform", no_update),
        tpl.get("kind", no_update),
    )


# --- KB browser search/filter (I) -------------------------------------------

@app.callback(
    Output("am-kb-faqs-pane",    "children"),
    Output("am-kb-replies-pane", "children"),
    Output("am-kb-search",    "value"),
    Output("am-kb-section",   "value"),
    Output("am-kb-platform",  "value"),
    Output("am-kb-sentiment", "value"),
    Input("am-kb-search-btn", "n_clicks"),
    Input("am-kb-reset-btn",  "n_clicks"),
    Input("am-kb-search",     "n_submit"),     # pressing Enter in the search box
    Input("am-kb-section",    "value"),
    Input("am-kb-platform",   "value"),
    Input("am-kb-sentiment",  "value"),
    State("am-kb-search",     "value"),
    prevent_initial_call=True,
)
def _am_kb_filter(_search_clicks, _reset_clicks, _enter_n,
                  sections, platforms, sentiments, query):
    """Filter the KB browser. Triggered by:
    - Click on the Search button
    - Pressing Enter in the search box
    - Click on the Clear button (resets all inputs)
    - Change in any of the three dropdowns
    """
    trig = ctx.triggered_id
    try:
        kb = _am_get_kb()
    except Exception as e:
        msg = dbc.Alert(f"KB load failed: {e}", color="danger")
        return msg, msg, no_update, no_update, no_update, no_update

    # Clear button: wipe all filter state and re-render with the full corpus.
    if trig == "am-kb-reset-btn":
        return (_am_render_faqs(kb.faqs),
                _am_render_replies(kb.comment_examples),
                "", None, None, None)

    q = (query or "").strip().lower()
    section_set   = set(sections or [])
    platform_set  = set(platforms or [])
    sentiment_set = set(sentiments or [])

    def _faq_match(f) -> bool:
        if section_set and (f.section or "") not in section_set:
            return False
        if not q:
            return True
        hay = " ".join([f.question, f.answer_short, f.answer_long, f.section]).lower()
        return q in hay

    def _reply_match(c) -> bool:
        if platform_set and (c.platform or "") not in platform_set:
            return False
        if sentiment_set and (c.sentiment or "") not in sentiment_set:
            return False
        if not q:
            return True
        return q in (c.comment + " " + c.response).lower()

    faqs    = [f for f in kb.faqs if _faq_match(f)]
    replies = [c for c in kb.comment_examples if _reply_match(c)]
    return (_am_render_faqs(faqs), _am_render_replies(replies),
            no_update, no_update, no_update, no_update)


# --- Delete past-reply exemplars from the KB browser ------------------------

@app.callback(
    Output("am-kb-replies-pane",   "children", allow_duplicate=True),
    Output("am-kb-delete-status",  "children"),
    Input({"type": "am-kb-delete", "row_id": dash.ALL}, "submit_n_clicks"),
    State("am-kb-platform",  "value"),
    State("am-kb-sentiment", "value"),
    State("am-kb-search",    "value"),
    prevent_initial_call=True,
)
def _am_kb_delete(_submits, platforms, sentiments, query):
    """Triggered by ConfirmDialogProvider for any past-reply row.
    Deletes the row from meta_comments.xlsx, refreshes the KB cache, and
    re-renders the past-replies pane honouring whatever filters are active.
    """
    triggered = ctx.triggered_id
    # Guard against the spurious initial fire (all submit_n_clicks None).
    if not triggered or not _submits or not any(s for s in _submits if s):
        return no_update, no_update
    if not isinstance(triggered, dict):
        return no_update, no_update
    row_id = triggered.get("row_id")
    if not row_id:
        return no_update, dbc.Alert("Cannot delete: row_id missing.",
                                     color="danger", className="mb-0 py-2 small")

    result = delete_exemplar(row_id)
    if not result.get("ok"):
        return no_update, dbc.Alert(
            f"⚠ Delete failed: {result.get('error', 'unknown error')}",
            color="danger", className="mb-0 py-2 small",
        )

    # Refresh the module-level KB cache so the next draft sees the deletion.
    with _am_kb_lock:
        _am_kb_cache["kb"] = result["kb"]
    kb = result["kb"]

    # Honour any active platform/sentiment/search filters in the re-render.
    q = (query or "").strip().lower()
    platform_set  = set(platforms or [])
    sentiment_set = set(sentiments or [])

    def _match(c) -> bool:
        if platform_set and (c.platform or "") not in platform_set:
            return False
        if sentiment_set and (c.sentiment or "") not in sentiment_set:
            return False
        if not q:
            return True
        return q in (c.comment + " " + c.response).lower()

    filtered = [c for c in kb.comment_examples if _match(c)]
    n = len(kb.comment_examples)
    toast = dbc.Alert([
        html.Strong("✓ Exemplar deleted"),
        html.Br(),
        html.Small(f"{n} past replies remain in the KB.", className="text-muted"),
    ], color="success", className="mb-0 py-2 small")
    return _am_render_replies(filtered), toast


# ===========================================================================
# ARCHIVE PAGE — every generated newsletter / blog, browsable
# ===========================================================================

_AR_GRADE_COLORS = {"pass": "success", "warn": "warning", "fail": "danger"}
_AR_KIND_ICON   = {"blog": "📝", "newsletter": "✉"}
_AR_FORMAT_META = _FORMAT_META   # reuse Create page's per-format colour map


def _ar_quality_badge(score, grade):
    """Inline quality-score badge (0-100). Hidden when no score."""
    if score is None:
        return html.Span("—", className="text-muted")
    color = ("success" if score >= 80
             else "warning" if score >= 60 else "danger")
    return dbc.Badge(f"{score}/100",
                     color=color, className="me-1",
                     style={"fontSize": "0.7rem"},
                     title=grade or "")


def _ar_compliance_badge(grade, pass_rate):
    if not grade:
        return html.Span("—", className="text-muted")
    color = _AR_GRADE_COLORS.get(grade, "secondary")
    label = grade.upper()
    if pass_rate is not None:
        label += f"  {int(pass_rate*100)}%"
    return dbc.Badge(label, color=color, style={"fontSize": "0.7rem"})


def _ar_format_pill_row(paths: dict) -> html.Span:
    """Inline coloured pills for each format on the row, each a download link."""
    bits = []
    for ext in _FORMAT_ORDER:
        if ext not in paths:
            continue
        label, color, tip = _AR_FORMAT_META[ext]
        fname = os.path.basename(paths[ext])
        bits.append(
            html.A(label.split(" ", 1)[0],   # just the icon, not the word
                   href=f"/downloads/{fname}",
                   download=fname,
                   title=f"{label} — {tip}",
                   className=f"badge bg-{color} me-1 text-decoration-none",
                   style={"fontSize": "0.65rem", "padding": "4px 6px"})
        )
    return html.Span(bits)


def _archive_page():
    cfg_local = cfg
    entries = list_archive(cfg_local.output_dir)
    stats = archive_stats(entries)

    # ---- Stats strip ---------------------------------------------------
    stat_row = dbc.Row([
        dbc.Col(_stat_card("Total generations",  str(stats["total"]),  "primary"),  md=3),
        dbc.Col(_stat_card("Blog posts",
                            str(stats["by_kind"].get("blog", 0)), "info"), md=3),
        dbc.Col(_stat_card("Newsletters",
                            str(stats["by_kind"].get("newsletter", 0)), "info"), md=3),
        dbc.Col(_stat_card(
            "Avg compliance pass-rate",
            (f"{int(stats['avg_compliance_pass_rate']*100)}%"
             if stats["avg_compliance_pass_rate"] is not None else "—"),
            "success",
        ), md=3),
    ], className="mb-3 g-3")

    # ---- Filter strip --------------------------------------------------
    filter_card = dbc.Card(dbc.CardBody([
        html.Div("Filter", className="section-eyebrow"),
        dbc.Row([
            dbc.Col([
                dbc.Label("Search title"),
                dbc.Input(id="ar-search", type="text",
                          placeholder="e.g. ISA, pension, FCA…",
                          debounce=True, n_submit=0),
            ], md=5),
            dbc.Col([
                dbc.Label("Kind"),
                dcc.Dropdown(id="ar-kind",
                             options=[
                                 {"label": "All",         "value": "all"},
                                 {"label": "Blog posts",  "value": "blog"},
                                 {"label": "Newsletters", "value": "newsletter"},
                             ],
                             value="all", clearable=False),
            ], md=2),
            dbc.Col([
                dbc.Label("From date"),
                dbc.Input(id="ar-from", type="date"),
            ], md=2),
            dbc.Col([
                dbc.Label("To date"),
                dbc.Input(id="ar-to",   type="date"),
            ], md=2),
            dbc.Col([
                dbc.Button("✕  Clear", id="ar-reset", color="light",
                            outline=True, size="sm",
                            style={"marginTop": "30px"}),
            ], md=1),
        ], className="g-3"),
    ]), className="mb-3")

    # ---- Initial table render -----------------------------------------
    table = _ar_render_table(entries)

    return html.Div([
        _page_header("Archive",
                     "Every generated newsletter and blog, browsable, "
                     "downloadable, and previewable in place."),
        stat_row,
        filter_card,
        # Toast slot for the delete callback
        html.Div(id="ar-status", className="mb-2"),
        # Table + preview pane in their own divs so search and click can
        # update them independently without re-rendering the filter strip.
        html.Div(table, id="ar-table-pane"),
        html.Div(id="ar-preview-pane", className="mt-3"),
    ])


def _ar_render_table(entries: list) -> html.Div:
    if not entries:
        return dbc.Alert(
            "No matching generations. Adjust the filters or generate something on the Create page.",
            color="info", className="mb-0",
        )
    rows = []
    for e in entries:
        rows.append(html.Tr([
            html.Td(_AR_KIND_ICON.get(e.kind, "•"),
                    style={"width": "30px", "textAlign": "center"}),
            html.Td([
                html.Div(e.title, className="fw-semibold",
                         style={"fontSize": "13px"}),
                html.Small(e.basename, className="text-muted",
                           style={"fontSize": "11px"}),
            ]),
            html.Td(html.Small(e.date, className="text-muted"),
                    style={"width": "100px"}),
            html.Td(_ar_compliance_badge(e.compliance_grade, e.compliance_pass_rate),
                    style={"width": "120px", "textAlign": "center"}),
            html.Td(_ar_quality_badge(e.quality_score, e.quality_grade),
                    style={"width": "90px",  "textAlign": "center"}),
            html.Td(_ar_format_pill_row(e.paths),
                    style={"width": "200px"}),
            html.Td([
                dbc.Button("Preview",
                           id={"type": "ar-preview-btn", "basename": e.basename},
                           color="primary", size="sm", outline=True,
                           className="me-1", n_clicks=0),
                dcc.ConfirmDialogProvider(
                    children=dbc.Button(
                        "🗑", color="danger", size="sm", outline=True,
                        title="Delete this generation (all formats)",
                    ),
                    id={"type": "ar-delete", "basename": e.basename},
                    message=(f"Delete every file for '{e.basename}'?\n\n"
                             "This removes the .html, .pdf, .docx, .md, .eml, "
                             ".txt and .json — and cannot be undone."),
                ),
            ], style={"width": "150px", "textAlign": "right"}),
        ]))
    return dbc.Card(dbc.CardBody([
        html.Div(f"{len(entries)} generation(s)",
                 className="section-eyebrow mb-2"),
        dbc.Table(
            [html.Thead(html.Tr([
                html.Th(""),
                html.Th("Title"),
                html.Th("Date"),
                html.Th("Compliance", style={"textAlign": "center"}),
                html.Th("Quality",    style={"textAlign": "center"}),
                html.Th("Formats"),
                html.Th("",            style={"textAlign": "right"}),
            ])),
             html.Tbody(rows)],
            hover=True, size="sm", striped=False,
            className="mb-0 align-middle",
        ),
    ]))


# --- Filter callback --------------------------------------------------------

@app.callback(
    Output("ar-table-pane",   "children"),
    Output("ar-search",       "value"),
    Output("ar-kind",         "value"),
    Output("ar-from",         "value"),
    Output("ar-to",           "value"),
    Input("ar-search",        "value"),
    Input("ar-search",        "n_submit"),
    Input("ar-kind",          "value"),
    Input("ar-from",          "value"),
    Input("ar-to",            "value"),
    Input("ar-reset",         "n_clicks"),
    prevent_initial_call=True,
)
def _ar_filter(query, _enter, kind, date_from, date_to, _reset):
    trig = ctx.triggered_id
    if trig == "ar-reset":
        # Wipe filters and re-render full archive.
        entries = list_archive(cfg.output_dir)
        return _ar_render_table(entries), "", "all", None, None
    entries = list_archive(
        cfg.output_dir,
        kind=(kind if kind and kind != "all" else None),
        query=query or None,
        date_from=date_from or None,
        date_to=date_to or None,
    )
    return (_ar_render_table(entries),
            no_update, no_update, no_update, no_update)


# --- Preview a single entry inline -----------------------------------------

@app.callback(
    Output("ar-preview-pane", "children"),
    Input({"type": "ar-preview-btn", "basename": dash.ALL}, "n_clicks"),
    prevent_initial_call=True,
)
def _ar_preview(_clicks):
    triggered = ctx.triggered_id
    if not triggered or not _clicks or not any(c for c in _clicks if c):
        return no_update
    if not isinstance(triggered, dict):
        return no_update
    basename = triggered.get("basename")
    if not basename:
        return no_update

    entry = archive_get_entry(cfg.output_dir, basename)
    if not entry:
        return dbc.Alert(f"Entry '{basename}' not found.",
                          color="warning", className="mb-0")

    # Read HTML for the iframe.
    html_path = entry.paths.get("html")
    preview_html = ""
    if html_path:
        try:
            with open(html_path) as f:
                preview_html = f.read()
        except OSError:
            preview_html = "<p>Could not load HTML preview.</p>"

    # Load full quality dict from JSON sidecar for the expandable report.
    quality_full = None
    if entry.json_path and os.path.isfile(entry.json_path):
        try:
            with open(entry.json_path, encoding="utf-8") as f:
                quality_full = json.load(f).get("quality") or None
        except (OSError, json.JSONDecodeError):
            pass

    download_buttons = []
    for ext in _FORMAT_ORDER:
        p = entry.paths.get(ext)
        if not p:
            continue
        label, color, tip = _AR_FORMAT_META[ext]
        fname = os.path.basename(p)
        download_buttons.append(dbc.Button(
            label, href=f"/downloads/{fname}", external_link=True,
            download=fname, color=color, size="sm",
            className="me-2 mb-2", title=tip,
        ))

    info_rows = [
        ("Kind",         entry.kind.title()),
        ("Generated at", entry.generated_at[:19].replace("T", " ")
                          if entry.generated_at else "—"),
        ("Date / version", f"{entry.date} · v{entry.version}"),
        ("Input articles", str(entry.input_article_count)),
        ("Sections",     str(entry.sections_count)),
    ]
    if entry.kind == "blog":
        info_rows += [
            ("Word count",    str(entry.word_count or "—")),
            ("Sources cited", str(entry.sources_cited_count)),
            ("Quality score", _ar_quality_badge(entry.quality_score,
                                                 entry.quality_grade)),
        ]
    info_rows.append((
        "Compliance",
        (f"{(entry.compliance_grade or '?').upper()}  "
         f"({int((entry.compliance_pass_rate or 0)*100)}% pass)"
         if entry.compliance_grade else "—"),
    ))

    info_table = dbc.Table([
        html.Tbody([
            html.Tr([
                html.Td(label, className="text-muted small",
                        style={"width": "150px"}),
                html.Td(value, style={"fontSize": "13px"}),
            ]) for label, value in info_rows
        ])
    ], borderless=True, size="sm", className="mb-0")

    # Expandable quality report — shown only for blogs with a quality block.
    quality_widget = html.Div()
    if entry.kind == "blog" and quality_full:
        quality_widget = dbc.Accordion([
            dbc.AccordionItem(
                _quality_card(quality_full),
                title=f"📊 Quality Report — {entry.quality_score}/100  ({entry.quality_grade})",
            ),
        ], start_collapsed=True, flush=False,
           className="mt-3 shadow-sm")

    return dbc.Card(dbc.CardBody([
        dbc.Row([
            dbc.Col([
                html.Div([
                    html.Span(_AR_KIND_ICON.get(entry.kind, "•"),
                              style={"fontSize": "1.4rem", "marginRight": "8px"}),
                    html.Strong(entry.title, style={"fontSize": "16px"}),
                ]),
                html.Small(entry.basename, className="text-muted"),
            ], md=10),
            dbc.Col([
                dbc.Button("✕", id="ar-preview-close",
                            color="light", size="sm",
                            className="float-end", n_clicks=0),
                dbc.Tooltip("Close the preview", target="ar-preview-close",
                            placement="left"),
            ], md=2),
        ], className="mb-3"),
        html.Div("Download as:", className="text-muted small mb-1"),
        html.Div(download_buttons, className="mb-3"),
        dbc.Row([
            dbc.Col([
                info_table,
                quality_widget,
            ], md=4),
            dbc.Col(html.Iframe(
                srcDoc=preview_html,
                style={"width": "100%", "height": "640px",
                       "border": "1px solid var(--warren-border)",
                       "borderRadius": "var(--radius)",
                       "background": "#ffffff"},
            ), md=8),
        ]),
    ]), className="shadow-sm",
        style={"borderLeft": "4px solid var(--bs-primary)"})


@app.callback(
    Output("ar-preview-pane", "children", allow_duplicate=True),
    Input("ar-preview-close", "n_clicks"),
    prevent_initial_call=True,
)
def _ar_preview_close(_n):
    return None if _n else no_update


# --- Delete an archived generation -----------------------------------------

@app.callback(
    Output("ar-table-pane", "children", allow_duplicate=True),
    Output("ar-status",     "children"),
    Output("ar-preview-pane","children", allow_duplicate=True),
    Input({"type": "ar-delete", "basename": dash.ALL}, "submit_n_clicks"),
    State("ar-search", "value"),
    State("ar-kind",   "value"),
    State("ar-from",   "value"),
    State("ar-to",     "value"),
    prevent_initial_call=True,
)
def _ar_delete(_submits, query, kind, date_from, date_to):
    triggered = ctx.triggered_id
    if not triggered or not _submits or not any(s for s in _submits if s):
        return no_update, no_update, no_update
    if not isinstance(triggered, dict):
        return no_update, no_update, no_update
    basename = triggered.get("basename")
    if not basename:
        return no_update, no_update, no_update

    result = archive_delete_entry(cfg.output_dir, basename)
    if not result.get("ok"):
        return (no_update,
                dbc.Alert(f"⚠ Delete failed: {result.get('error', 'unknown')}",
                          color="danger", className="mb-0 py-2 small"),
                no_update)

    # Re-render with current filters honoured.
    entries = list_archive(
        cfg.output_dir,
        kind=(kind if kind and kind != "all" else None),
        query=query or None,
        date_from=date_from or None,
        date_to=date_to or None,
    )
    toast = dbc.Alert([
        html.Strong("✓ Deleted "), html.Code(basename),
        html.Small(f"  ({result['count']} files removed)",
                   className="text-muted ms-2"),
    ], color="success", className="mb-0 py-2 small")
    # Close the preview pane in case it was showing the just-deleted entry.
    return _ar_render_table(entries), toast, None


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    app.run(debug=False, port=8050)
